#!/usr/bin/env python3
"""
SAP Cloud Integration iFlow - Generatore documentazione v4
Migliorie: tabella CITS, colonna Transform, Resource Inventory Used/Not used
sadsa
Uso: pip install python-docx && python scpi_iflow_doc_generator.py
"""
from __future__ import annotations
import hashlib
import json
import re
import sys
import zipfile
import os
import xml.etree.ElementTree as ET
from datetime import datetime, timezone
from pathlib import Path
from typing import Any
try:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor
    from io import BytesIO
    from PIL import Image, ImageDraw, ImageFont
except ImportError:
    print("pip install python-docx")
    sys.exit(1)

# ── Config ───────────────────────────────────────────────────────────────
INPUT_DIR = Path("input")
OUTPUT_DIR = Path("output")
DOCS_DIR = OUTPUT_DIR / "docs"
JSON_DIR = OUTPUT_DIR / "json"
MASK_SECRETS = True
SENSITIVE_KEYS = (
    "password",
    "passwd",
    "secret",
    "token",
    "credential",
    "authorization",
    "apikey",
    "api_key",
    "clientsecret",
    "privatekey",
)

# ── Utility ──────────────────────────────────────────────────────────────
def local_name(tag: str) -> str:
    return tag.split("}", 1)[-1] if "}" in tag else tag

def clean_text(v: Any) -> str:
    if v is None:
        return ""
    return re.sub(r"\s+", " ", str(v)).strip()

def safe_filename(v: str) -> str:
    v = clean_text(v)
    v = re.sub(r'[<>:"/\\|?*]+', "_", v)
    return v[:180] or "unnamed"

def now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()

def sha256_bytes(b: bytes) -> str:
    return hashlib.sha256(b).hexdigest()

def is_sensitive(k: str) -> bool:
    k = clean_text(k).lower().replace("-", "").replace("_", "")
    return any(w.replace("-", "").replace("_", "") in k for w in SENSITIVE_KEYS)

def mask_value(k: str, v: Any) -> str:
    t = clean_text(v)
    if not MASK_SECRETS:
        return t
    if is_sensitive(k):
        return "<configurato>"
    if re.match(r"^(Bearer\s+|Basic\s+)", t, re.IGNORECASE):
        return "<configurato>"
    return t

def unique_list(items: list[str]) -> list[str]:
    seen = set()
    result = []
    for item in items:
        item = clean_text(item)
        if item and item not in seen:
            seen.add(item)
            result.append(item)
    return result

# ── ZIP / MANIFEST ───────────────────────────────────────────────────────
def parse_manifest(text: str) -> dict[str, str]:
    unfolded = []
    manifest = {}
    for line in text.splitlines():
        if line.startswith(" ") and unfolded:
            unfolded[-1] += line[1:]
        else:
            unfolded.append(line)
    for line in unfolded:
        if ": " in line:
            key, value = line.split(": ", 1)
            manifest[clean_text(key)] = clean_text(value)
    return manifest

def find_iflow_file(names: list[str]) -> str | None:
    candidates = [name for name in names if name.lower().endswith(".iflw")]
    if not candidates:
        return None
    candidates.sort(
        key=lambda name: (
            0 if "integrationflow" in name.lower() else 1,
            len(name),
        )
    )
    return candidates[0]

# ── BPMN ─────────────────────────────────────────────────────────────────
def get_element_properties(el: ET.Element) -> dict[str, str]:
    result = {}
    for node in el.iter():
        if local_name(node.tag).lower() != "property":
            continue
        key = clean_text(
            node.attrib.get("key")
            or node.attrib.get("name")
            or node.attrib.get("propertyName")
        )
        value = clean_text(node.attrib.get("value") or node.attrib.get("propertyValue"))
        for child in list(node):
            child_name = local_name(child.tag).lower()
            if child_name in ("key", "name", "propertyname") and not key:
                key = clean_text(child.text)
            elif child_name in ("value", "propertyvalue") and not value:
                value = clean_text(child.text)
        if key:
            result[key] = mask_value(key, value)
    return result

def get_direct_child_text(el: ET.Element, child_name: str) -> str:
    for child in list(el):
        if local_name(child.tag) == child_name:
            return clean_text(child.text)
    return ""

def parse_participants(root: ET.Element) -> list[dict]:
    participants = []
    for el in root.iter():
        if local_name(el.tag) != "participant":
            continue
        participants.append(
            {
                "id": clean_text(el.attrib.get("id")),
                "name": clean_text(el.attrib.get("name")),
                "type": clean_text(
                    el.attrib.get("participantType") or el.attrib.get("type")
                ),
                "process_ref": clean_text(
                    el.attrib.get("processRef") or el.attrib.get("processref")
                ),
                "properties": get_element_properties(el),
            }
        )
    return participants

def parse_message_flows(root: ET.Element) -> list[dict]:
    flows = []
    for el in root.iter():
        if local_name(el.tag) != "messageFlow":
            continue
        props = get_element_properties(el)
        adapter = (
            props.get("componentType")
            or props.get("Adapter Type")
            or props.get("adapterType")
            or props.get("Type")
            or ""
        )
        address = props.get("Address") or props.get("address") or props.get("URL") or ""
        flows.append(
            {
                "id": clean_text(el.attrib.get("id")),
                "name": clean_text(el.attrib.get("name")),
                "source_ref": clean_text(el.attrib.get("sourceRef")),
                "target_ref": clean_text(el.attrib.get("targetRef")),
                "adapter_type": clean_text(adapter),
                "address": clean_text(address),
                "properties": props,
            }
        )
    return flows

def normalize_key(value: str) -> str:
    return re.sub(r"[^a-z0-9]", "", clean_text(value).lower())

def get_step_component_type(
    bpmn_type: str,
    properties: dict[str, str],
    step_name: str = "",
) -> str:
    normalized_properties = {
        normalize_key(key): clean_text(value) for key, value in properties.items()
    }
    candidate_keys = (
        "componenttype",
        "activitytype",
        "steptype",
        "tasktype",
        "component",
        "adaptertype",
        "adapter",
    )
    searchable_text = " ".join(
        [
            clean_text(step_name),
            clean_text(bpmn_type),
            *normalized_properties.values(),
        ]
    ).lower()
    component_patterns = (
        ("content modifier", ("contentmodifier", "content modifier")),
        ("script groovy", ("groovy", "script")),
        ("message mapping", ("messagemapping", "message mapping")),
        ("trasformazione XSLT", ("xslt", "xsl transformation")),
        ("request reply", ("requestreply", "request reply")),
        ("invio messaggio", ("send", "sender")),
        ("ricezione messaggio", ("receive", "receiver")),
        ("router", ("router", "exclusivegateway", "choice")),
        ("splitter", ("splitter", "general splitter", "iterating splitter")),
        ("gather", ("gather", "aggregator")),
        ("multicast", ("multicast",)),
        ("content enricher", ("contentenricher", "content enricher")),
        ("process call", ("processcall", "process call")),
        ("idoc", ("idoc",)),
        ("odata", ("odata",)),
        ("soap", ("soap",)),
        ("http", ("http", "rest")),
        ("sftp", ("sftp",)),
        ("jms", ("jms",)),
        ("mail", ("mail",)),
    )
    for label, patterns in component_patterns:
        if any(pattern in searchable_text for pattern in patterns):
            return label
    if bpmn_type == "subProcess":
        return "sottoprocesso"
    if bpmn_type == "callActivity":
        for key in ("calledelement", "processref", "subprocess"):
            if normalized_properties.get(key):
                return "chiamata a un sottoprocesso"
        return "attività di integrazione CPI"
    return clean_text(bpmn_type)

def detect_mapping_expression(value: Any) -> dict:
    text = clean_text(value)
    if not text:
        return {"is_expression": False, "types": [], "matches": []}
    patterns = [
        ("Simple expression", r"\$\{[^}]+\}"),
        ("Parametro esternalizzato", r"\{\{[^}]+\}\}"),
        ("XPath/XML", r"(?<![\w$])(?:/[/\w:*@\[\]\-\.]+)"),
        ("JSONPath", r"(?<![\w])\$(?:\.[A-Za-z_][\w-]*|\[['\"][^'\"]+['\"]\]|\[\d+\])+"),
    ]
    detected_types = []
    matches = []
    for label, pattern in patterns:
        found = re.findall(pattern, text)
        if found:
            detected_types.append(label)
            matches.extend(found)
    return {
        "is_expression": bool(detected_types),
        "types": unique_list(detected_types),
        "matches": unique_list(matches),
    }

def extract_node_key_value(node: ET.Element) -> tuple[str, str]:
    key = clean_text(
        node.attrib.get("key")
        or node.attrib.get("name")
        or node.attrib.get("headerName")
        or node.attrib.get("headername")
        or node.attrib.get("propertyName")
        or node.attrib.get("propertyname")
        or node.attrib.get("id")
    )
    value = clean_text(
        node.attrib.get("value")
        or node.attrib.get("headerValue")
        or node.attrib.get("headervalue")
        or node.attrib.get("propertyValue")
        or node.attrib.get("propertyvalue")
        or node.attrib.get("content")
        or node.attrib.get("body")
    )
    for child in list(node):
        child_name = normalize_key(local_name(child.tag))
        if child_name in (
            "key", "name", "headername", "propertyname", "fieldname",
        ) and not key:
            key = clean_text(child.text)
        elif child_name in (
            "value", "headervalue", "propertyvalue", "content",
            "payload", "body", "messagebody", "messagecontent",
        ) and not value:
            value = clean_text(child.text)
    if not value and clean_text(node.text):
        value = clean_text(node.text)
    return key, value

def classify_content_modifier_entry(node_tag: str, key: str) -> str:
    normalized_tag = normalize_key(node_tag)
    normalized_key = normalize_key(key)
    body_markers = ("body", "payload", "messagebody", "messagecontent", "contentbody", "setbody")
    header_markers = ("header", "messageheader", "setheader", "customheader")
    property_markers = ("exchangeproperty", "messageproperty", "property", "setproperty", "exchange")
    if any(marker in normalized_tag for marker in body_markers) or any(marker in normalized_key for marker in body_markers):
        return "Body"
    if any(marker in normalized_tag for marker in header_markers) or any(marker in normalized_key for marker in header_markers):
        return "Header"
    if any(marker in normalized_tag for marker in property_markers) or any(marker in normalized_key for marker in property_markers):
        return "Exchange Property"
    return "Altro"
def get_node_text(node: ET.Element) -> str:
    """
    Recupera tutto il testo contenuto nel nodo e nei suoi discendenti.
    Utile per value CDATA, XML inline e body/payload annidati.
    """
    return clean_text("".join(node.itertext()))


def get_child_value(node: ET.Element, names: tuple[str, ...]) -> str:
    """
    Cerca il valore in attributi o figli diretti/annidati con uno dei nomi indicati.
    """
    normalized_names = {normalize_key(name) for name in names}

    for attr_name, attr_value in node.attrib.items():
        if normalize_key(attr_name) in normalized_names:
            return clean_text(attr_value)

    for child in node.iter():
        if child is node:
            continue

        child_name = normalize_key(local_name(child.tag))

        if child_name in normalized_names:
            value = clean_text(
                child.attrib.get("value")
                or child.attrib.get("content")
                or child.attrib.get("text")
                or get_node_text(child)
            )

            if value:
                return value

    return ""

def preserve_text(value: Any) -> str:
    """
    Mantiene i ritorni a capo del body, normalizzando solo CRLF/CR.
    """
    if value is None:
        return ""

    return str(value).replace("\r\n", "\n").replace("\r", "\n").strip()


def get_node_raw_text(node: ET.Element) -> str:
    """
    Estrae il testo completo di un nodo, inclusi CDATA e figli testuali,
    senza schiacciare gli spazi o i ritorni a capo.
    """
    return preserve_text("".join(node.itertext()))


def analyze_body_content(value: Any) -> dict:
    """
    Analizza il Body di un Content Modifier.

    Rileva:
    - formato XML / JSON / testo;
    - root XML e numero elementi;
    - chiavi JSON di primo livello;
    - espressioni CPI, XPath e JSONPath;
    - preview e lunghezza.
    """
    raw_value = preserve_text(value)

    result = {
        "raw_value": raw_value,
        "format": "Vuoto",
        "xml_root": "",
        "xml_elements": 0,
        "json_keys": [],
        "expression_types": [],
        "expression_matches": [],
        "has_mapping_expression": False,
        "length": len(raw_value),
        "preview": raw_value[:500],
    }

    if not raw_value:
        return result

    expression_info = detect_mapping_expression(raw_value)

    result["has_mapping_expression"] = expression_info["is_expression"]
    result["expression_types"] = expression_info["types"]
    result["expression_matches"] = expression_info["matches"]

    # JSON
    try:
        json_data = json.loads(raw_value)

        result["format"] = "JSON"

        if isinstance(json_data, dict):
            result["json_keys"] = list(json_data.keys())[:20]
        elif isinstance(json_data, list):
            result["json_keys"] = [
                f"Array ({len(json_data)} elementi)"
            ]

        return result

    except (json.JSONDecodeError, TypeError):
        pass

    # XML
    try:
        xml_root = ET.fromstring(raw_value)

        result["format"] = "XML"
        result["xml_root"] = local_name(xml_root.tag)
        result["xml_elements"] = sum(
            1
            for _ in xml_root.iter()
        )

        return result

    except ET.ParseError:
        pass

    # Template XML: può contenere ${...} o {{...}} e non essere XML valido.
    if "<" in raw_value and ">" in raw_value:
        result["format"] = "Template XML"

    elif result["has_mapping_expression"]:
        result["format"] = "Espressione / template CPI"

    else:
        result["format"] = "Testo"

    return result

CONTENT_MODIFIER_BODY_MARKERS = (
    "body",
    "payload",
    "messagebody",
    "messagecontent",
    "contentbody",
    "setbody",
)

CONTENT_MODIFIER_HEADER_MARKERS = (
    "header",
    "headers",
    "messageheader",
    "messageheaders",
    "setheader",
    "setheaders",
    "customheader",
    "customheaders",
)

CONTENT_MODIFIER_PROPERTY_MARKERS = (
    "exchangeproperty",
    "exchangeproperties",
    "messageproperty",
    "messageproperties",
    "setproperty",
    "setproperties",
    "properties",
)


def detect_category(
    node: ET.Element,
    inherited_category: str = "",
) -> str:
    """
    Determina Header, Exchange Property o Body dal nodo CPI corrente
    oppure eredita la categoria identificata nel contenitore padre.
    """
    node_name = normalize_key(local_name(node.tag))

    if any(
        marker in node_name
        for marker in CONTENT_MODIFIER_BODY_MARKERS
    ):
        return "Body"

    if any(
        marker in node_name
        for marker in CONTENT_MODIFIER_HEADER_MARKERS
    ):
        return "Header"

    if any(
        marker in node_name
        for marker in CONTENT_MODIFIER_PROPERTY_MARKERS
    ):
        return "Exchange Property"

    return inherited_category

def parse_content_modifier(el: ET.Element) -> dict:
    """
    Estrae Header, Exchange Properties e Body da Content Modifier SAP CPI.

    Supporta il formato CPI tipico:
    - headerTable: XML con <row><cell id="...">...</cell></row>;
    - propertyTable: XML con la medesima struttura;
    - wrapContent: body/payload JSON, XML o template;
    - bodyType: metadato del body (constant, expression ecc.).

    Rileva anche espressioni CPI:
    ${...}, {{...}}, XPath/XML e JSONPath.
    """
    entries = []
    seen = set()

    def preserve_text(value: Any) -> str:
        if value is None:
            return ""

        return (
            str(value)
            .replace("\r\n", "\n")
            .replace("\r", "\n")
            .strip()
        )

    def add_entry(
        category: str,
        name: str,
        value: str,
        value_type: str = "",
    ):
        """
        Inserisce una riga evitando duplicati.
        value_type è, ad esempio: constant, expression, xpath.
        """
        name = clean_text(name)

        if category == "Body":
            raw_value = preserve_text(value)
            normalized_value = raw_value

            if not name:
                name = "Body"
        else:
            raw_value = clean_text(value)
            normalized_value = raw_value

        if not name and not normalized_value:
            return

        unique_key = (
            category,
            name,
            normalized_value,
            clean_text(value_type),
        )

        if unique_key in seen:
            return

        seen.add(unique_key)

        expression_info = detect_mapping_expression(normalized_value)

        # Il type CPI "expression" è già un'indicazione di mapping,
        # anche se il valore non contiene una sintassi intercettata.
        is_expression = (
            expression_info["is_expression"]
            or clean_text(value_type).lower()
            in ("expression", "xpath", "jsonpath", "simple")
        )

        expression_types = list(expression_info["types"])

        if (
            clean_text(value_type)
            and clean_text(value_type) not in expression_types
        ):
            expression_types.insert(
                0,
                f"Tipo CPI: {clean_text(value_type)}",
            )

        entry = {
            "category": category,
            "name": name or "-",
            "value": mask_value(name or category, normalized_value) or "-",
            "value_type": clean_text(value_type),
            "is_expression": is_expression,
            "expression_types": unique_list(expression_types),
            "expression_matches": expression_info["matches"],
        }

        if category == "Body":
            body_analysis = {
                "raw_value": raw_value,
                "format": "Vuoto",
                "length": len(raw_value),
                "xml_root": "",
                "xml_elements": 0,
                "json_keys": [],
                "has_mapping_expression": is_expression,
                "expression_types": unique_list(expression_types),
                "expression_matches": expression_info["matches"],
                "preview": raw_value[:500],
            }

            if raw_value:
                try:
                    json_value = json.loads(raw_value)
                    body_analysis["format"] = "JSON"

                    if isinstance(json_value, dict):
                        body_analysis["json_keys"] = list(
                            json_value.keys()
                        )[:20]

                    elif isinstance(json_value, list):
                        body_analysis["json_keys"] = [
                            f"Array ({len(json_value)} elementi)"
                        ]

                except (json.JSONDecodeError, TypeError):
                    try:
                        xml_root = ET.fromstring(raw_value)
                        body_analysis["format"] = "XML"
                        body_analysis["xml_root"] = local_name(xml_root.tag)
                        body_analysis["xml_elements"] = sum(
                            1 for _ in xml_root.iter()
                        )

                    except ET.ParseError:
                        if "<" in raw_value and ">" in raw_value:
                            body_analysis["format"] = "Template XML"

                        elif is_expression:
                            body_analysis["format"] = (
                                "Espressione / template CPI"
                            )

                        else:
                            body_analysis["format"] = "Testo"

            entry["body_analysis"] = body_analysis

        entries.append(entry)

    def parse_cpi_table(raw_table: Any, category: str):
        """
        Legge headerTable/propertyTable CPI.

        Formato atteso:
        <row>
            <cell id="Action">Create</cell>
            <cell id="Type">constant</cell>
            <cell id="Value">...</cell>
            <cell id="Name">S4_path</cell>
        </row>
        """
        raw_table = preserve_text(raw_table)

        if not raw_table:
            return

        try:
            # headerTable/propertyTable spesso contengono più <row>
            # senza un root XML comune.
            root = ET.fromstring(f"<rows>{raw_table}</rows>")

        except ET.ParseError:
            return

        for row in root.iter():
            if local_name(row.tag).lower() != "row":
                continue

            columns = {}

            for cell in list(row):
                if local_name(cell.tag).lower() != "cell":
                    continue

                column_id = clean_text(
                    cell.attrib.get("id")
                    or cell.attrib.get("name")
                    or cell.attrib.get("key")
                ).lower()

                column_value = preserve_text("".join(cell.itertext()))

                if column_id:
                    columns[column_id] = column_value

            name = (
                columns.get("name")
                or columns.get("key")
                or columns.get("headername")
                or columns.get("propertyname")
                or ""
            )

            value = (
                columns.get("value")
                or columns.get("content")
                or columns.get("payload")
                or ""
            )

            value_type = (
                columns.get("type")
                or columns.get("valuetype")
                or ""
            )

            action = columns.get("action", "").lower()

            # Ignore righe di cancellazione senza valore/nome utile.
            if action in ("delete", "remove") and not name and not value:
                continue

            if name or value:
                add_entry(
                    category=category,
                    name=name,
                    value=value,
                    value_type=value_type,
                )

    def find_property_value(
        properties: dict[str, str],
        candidate_keys: tuple[str, ...],
    ) -> str:
        """
        Recupera una property ignorando maiuscole/minuscole e separatori.
        """
        normalized_properties = {
            normalize_key(key): value
            for key, value in properties.items()
        }

        for candidate in candidate_keys:
            value = normalized_properties.get(normalize_key(candidate))

            if value is not None:
                return preserve_text(value)

        return ""

    # Le extension property possono essere estratte anche da attributi
    # e XML annidato, a seconda della versione di export CPI.
    properties = get_element_properties(el)

    # Header configurati nella Header Table CPI.
    header_table = find_property_value(
        properties,
        (
            "headerTable",
            "headersTable",
            "messageHeaders",
            "setHeaders",
        ),
    )

    parse_cpi_table(
        raw_table=header_table,
        category="Header",
    )

    # Exchange Property configurate nella Property Table CPI.
    property_table = find_property_value(
        properties,
        (
            "propertyTable",
            "propertiesTable",
            "exchangePropertyTable",
            "exchangeProperties",
            "messageProperties",
            "setProperties",
        ),
    )

    parse_cpi_table(
        raw_table=property_table,
        category="Exchange Property",
    )

    # Body CPI: normalmente è wrapContent.
    body_value = find_property_value(
        properties,
        (
            "wrapContent",
            "bodyValue",
            "body",
            "messageBody",
            "messageContent",
            "payload",
            "content",
        ),
    )

    body_type = find_property_value(
        properties,
        (
            "bodyType",
            "bodyContentType",
            "payloadType",
            "contentType",
        ),
    )

    if body_value:
        add_entry(
            category="Body",
            name="Body",
            value=body_value,
            value_type=body_type,
        )

    # Fallback per formati CPI differenti: non aggiunge le property tecniche
    # headerTable/propertyTable/bodyType come fossero dati funzionali.
    for key, value in properties.items():
        normalized_key = normalize_key(key)

        if normalized_key in (
            "headertable",
            "propertiestable",
            "propertytable",
            "exchangepropertytable",
            "wrapcontent",
            "bodytype",
            "bodycontenttype",
            "payloadtype",
        ):
            continue

        if any(
            marker in normalized_key
            for marker in (
                "headername",
                "headervalue",
                "setheader",
            )
        ):
            add_entry("Header", key, value)

        elif any(
            marker in normalized_key
            for marker in (
                "propertyname",
                "propertyvalue",
                "setproperty",
                "exchangeproperty",
                "messageproperty",
            )
        ):
            add_entry("Exchange Property", key, value)

        elif any(
            marker in normalized_key
            for marker in (
                "bodyvalue",
                "messagebodyvalue",
                "payloadvalue",
            )
        ):
            add_entry(
                "Body",
                "Body",
                value,
                body_type,
            )

    headers = [
        entry
        for entry in entries
        if entry["category"] == "Header"
    ]

    exchange_properties = [
        entry
        for entry in entries
        if entry["category"] == "Exchange Property"
    ]

    bodies = [
        entry
        for entry in entries
        if entry["category"] == "Body"
    ]

    return {
        "headers": headers,
        "exchange_properties": exchange_properties,
        "bodies": bodies,
        "entries": entries,
        "has_mapping_expression": any(
            entry["is_expression"]
            for entry in entries
        ),
    }  

def parse_processes(root: ET.Element) -> list[dict]:
    processes = []
    ignored = {"extensionElements", "documentation", "incoming", "outgoing", "laneSet", "lane"}
    for proc in root.iter():
        if local_name(proc.tag) != "process":
            continue
        process_data = {
            "id": clean_text(proc.attrib.get("id")),
            "name": clean_text(proc.attrib.get("name")),
            "properties": get_element_properties(proc),
            "steps": [],
            "connections": [],
        }
        for el in list(proc):
            element_type = local_name(el.tag)
            if element_type in ignored:
                continue
            if element_type == "sequenceFlow":
                process_data["connections"].append({
                    "id": clean_text(el.attrib.get("id")),
                    "name": clean_text(el.attrib.get("name")),
                    "source_ref": clean_text(el.attrib.get("sourceRef")),
                    "target_ref": clean_text(el.attrib.get("targetRef")),
                    "condition": get_direct_child_text(el, "conditionExpression"),
                })
                continue
            if not el.attrib.get("id"):
                continue
            props = get_element_properties(el)

            step_name = clean_text(el.attrib.get("name"))

            component_type = get_step_component_type(
                bpmn_type=clean_text(element_type),
                properties=props,
                step_name=step_name,
            )

            # Analizza lo step anche se CPI lo espone come callActivity.
            content_modifier = parse_content_modifier(el)

            # Se sono trovati Header, Property o Body,
            # lo step viene classificato come Content Modifier.
            if content_modifier["entries"]:
                component_type = "content modifier"
            else:
                content_modifier = None

            step = {
                "id": clean_text(el.attrib.get("id")),
                "name": clean_text(el.attrib.get("name")),
                "bpmn_type": clean_text(element_type),
                "called_element": clean_text(
                    el.attrib.get("calledElement") or el.attrib.get("calledelement")
                ),
                "documentation": get_direct_child_text(el, "documentation"),
                "properties": props,
                "component_type": component_type,
                "content_modifier": content_modifier,
                "incoming": [],
                "outgoing": [],
            }
            for child in list(el):
                child_name = local_name(child.tag)
                if child_name == "incoming":
                    step["incoming"].append(clean_text(child.text))
                elif child_name == "outgoing":
                    step["outgoing"].append(clean_text(child.text))
            process_data["steps"].append(step)
        processes.append(process_data)
    return processes

def order_process_steps(process: dict) -> list[dict]:
    steps = process.get("steps", [])
    connections = process.get("connections", [])
    if not steps:
        return []
    steps_by_id = {step["id"]: step for step in steps if step.get("id")}
    outgoing: dict[str, list[str]] = {step_id: [] for step_id in steps_by_id}
    incoming_count = {step_id: 0 for step_id in steps_by_id}
    for connection in connections:
        source_id = connection.get("source_ref")
        target_id = connection.get("target_ref")
        if source_id in steps_by_id and target_id in steps_by_id:
            outgoing[source_id].append(target_id)
            incoming_count[target_id] += 1
    start_ids = [
        step["id"]
        for step in steps
        if step.get("bpmn_type") == "startEvent" and step.get("id") in steps_by_id
    ]
    no_incoming_ids = [
        step["id"]
        for step in steps
        if step.get("id") in steps_by_id
        and incoming_count[step["id"]] == 0
        and step["id"] not in start_ids
    ]
    queue = start_ids + no_incoming_ids
    if not queue:
        queue = [step["id"] for step in steps if step.get("id")]
    ordered = []
    visited = set()
    while queue:
        current_id = queue.pop(0)
        if current_id in visited:
            continue
        visited.add(current_id)
        ordered.append(steps_by_id[current_id])
        for target_id in outgoing.get(current_id, []):
            incoming_count[target_id] -= 1
            if incoming_count[target_id] <= 0 and target_id not in visited:
                queue.append(target_id)
    for step in steps:
        if step.get("id") and step["id"] not in visited:
            ordered.append(step)
    return ordered

def normalize_process_name(value: str) -> str:
    return re.sub(r"[^a-z0-9]", "", clean_text(value).lower())

def order_processes(processes: list[dict]) -> list[dict]:
    if not processes:
        return []
    processes_by_id = {process["id"]: process for process in processes if process.get("id")}
    processes_by_name = {
        normalize_process_name(process.get("name") or process.get("id")): process
        for process in processes
    }
    integration_process = next(
        (
            process
            for process in processes
            if normalize_process_name(process.get("name")) == "integrationprocess"
        ),
        processes[0],
    )
    ordered = []
    added_ids = set()

    def add_process(process: dict):
        process_id = process.get("id")
        if not process_id or process_id in added_ids:
            return
        added_ids.add(process_id)
        ordered.append(process)
        for step in order_process_steps(process):
            if step.get("bpmn_type") != "callActivity":
                continue
            called_element = clean_text(step.get("called_element"))
            called_process = None
            if called_element:
                called_process = processes_by_id.get(called_element)
                if not called_process:
                    called_process = processes_by_name.get(normalize_process_name(called_element))
            if not called_process:
                called_process = processes_by_name.get(normalize_process_name(step.get("name")))
            if called_process:
                add_process(called_process)

    add_process(integration_process)
    for process in processes:
        add_process(process)
    return ordered

def find_router_details(processes: list[dict]) -> list[dict]:
    routers = []
    for process in processes:
        connections_by_id = {
            connection["id"]: connection
            for connection in process["connections"]
            if connection.get("id")
        }
        for step in process["steps"]:
            if step["bpmn_type"] not in ("exclusiveGateway", "inclusiveGateway"):
                continue
            routes = []
            for outgoing_id in step["outgoing"]:
                connection = connections_by_id.get(outgoing_id, {})
                routes.append({
                    "sequence_flow_id": outgoing_id,
                    "name": connection.get("name", ""),
                    "condition": connection.get("condition", "") or "Default",
                    "target_ref": connection.get("target_ref", ""),
                })
            routers.append({
                "process_id": process["id"],
                "process_name": process["name"],
                "id": step["id"],
                "name": step["name"],
                "type": step["bpmn_type"],
                "routes": routes,
            })
    return routers

# ── Parametri ────────────────────────────────────────────────────────────
def parse_externalized_parameters(files: dict[str, bytes]) -> list[dict]:
    params = {}
    for name in files:
        if name.lower().endswith("parameters.propdef"):
            content = files[name].decode("utf-8", errors="replace")
            for line in content.splitlines():
                line = clean_text(line)
                if not line or line.startswith("#") or "=" not in line:
                    continue
                key, value = line.split("=", 1)
                key = clean_text(key)
                params.setdefault(key, {"name": key, "type": "", "value": "", "source": name})
                params[key]["type"] = clean_text(value)
    for name in files:
        if name.lower().endswith("parameters.prop"):
            content = files[name].decode("utf-8", errors="replace")
            for line in content.splitlines():
                line = clean_text(line)
                if not line or line.startswith("#") or "=" not in line:
                    continue
                key, value = line.split("=", 1)
                key = clean_text(key)
                params.setdefault(key, {"name": key, "type": "", "value": "", "source": name})
                params[key]["value"] = mask_value(key, value)
    return sorted(params.values(), key=lambda item: item["name"].lower())

# ── Risorse ──────────────────────────────────────────────────────────────
def parse_resources(files: dict[str, bytes]) -> list[dict]:
    resources = []
    resource_types = {
        ".groovy": "Groovy Script",
        ".mmap": "Message Mapping",
        ".xsd": "XSD",
        ".wsdl": "WSDL",
        ".json": "JSON",
        ".xml": "XML",
        ".csv": "CSV",
        ".xslt": "XSLT",
        ".xsl": "XSLT",
    }
    for name, content in files.items():
        suffix = Path(name).suffix.lower()
        if suffix in resource_types:
            resources.append({
                "path": name,
                "name": Path(name).name,
                "type": resource_types[suffix],
                "size_bytes": len(content),
                "sha256": sha256_bytes(content),
            })
    return sorted(resources, key=lambda item: item["path"].lower())

# ── XSLT parser ──────────────────────────────────────────────────────────
XSL_NS = "http://www.w3.org/1999/XSL/Transform"

def parse_xslt(files: dict[str, bytes]) -> list[dict]:
    results = []
    for name, content in files.items():
        if Path(name).suffix.lower() not in (".xslt", ".xsl"):
            continue
        try:
            root = ET.fromstring(content)
        except ET.ParseError:
            results.append({"name": Path(name).name, "path": name, "classifications": ["errore_parsing"]})
            continue
        categories = classify_xslt(root, Path(name).stem.lower())
        results.append({"name": Path(name).name, "path": name, "classifications": categories})
    return results

def classify_xslt(root: ET.Element, stem: str) -> list[str]:
    categories = []
    for template in root.iter(f"{{{XSL_NS}}}template"):
        if template.attrib.get("match", "") in ("/", "/*"):
            for child in template:
                if local_name(child.tag) == "element":
                    categories.append("cambio_root_tag")
                    break
    if "removemultimap" in stem or "remove" in stem:
        categories.append("rimozione_struttura")
    if "changeroottag" in stem:
        categories.append("cambio_root_tag")
    if (
        root.find(f".//{{{XSL_NS}}}value-of") is not None
        and root.find(f".//{{{XSL_NS}}}template") is not None
        and "cambio_root_tag" not in categories
    ):
        categories.append("mapping")
    if root.find(f".//{{{XSL_NS}}}if") is not None or root.find(f".//{{{XSL_NS}}}choose") is not None:
        categories.append("filtro")
    for element in root.iter(f"{{{XSL_NS}}}element"):
        for child in element:
            if local_name(child.tag) == "text" and clean_text(child.text):
                categories.append("arricchimento")
                break
    return unique_list(categories or ["trasformazione_generica"])

# ── Groovy parser ────────────────────────────────────────────────────────
PURPOSE_PATTERNS = {
    "buildRows": "Costruisce le righe strutturate del payload, separando taglie e colori per l'invio a SAP",
    "checkProperty": "Verifica che un campo esista e sia valorizzato prima di proseguire l'elaborazione",
    "checkUOM": "Controlla che le unità di misura ricevute siano valide e coerenti con quelle attese da SAP",
    "convertStringToBoolean": "Converte valori testuali (true/false) nel formato booleano richiesto da SAP",
    "mapFiberCode": "Traduce i codici fibra da Centric ai codici standard SAP",
    "removeFieldsUpdate": "In caso di aggiornamento, rimuove i campi che non devono essere modificati per evitare errori su SAP",
}

def guess_script_purpose(name: str, functions: list[str], strings: list[str]) -> str:
    script_name = Path(name).stem.lower()
    for pattern, description in PURPOSE_PATTERNS.items():
        if pattern.lower() in script_name:
            return description
    for function_name in functions:
        function_name_lower = function_name.lower()
        if "valid" in function_name_lower or "check" in function_name_lower:
            return "Esegue controlli e validazioni sui dati del messaggio"
        if "build" in function_name_lower or "create" in function_name_lower:
            return "Costruisce strutture dati da inviare al sistema destinatario"
        if "map" in function_name_lower or "convert" in function_name_lower:
            return "Converte o mappa dati tra formati diversi"
    combined = " ".join(strings).lower()
    if "error" in combined or "exception" in combined:
        return "Gestisce condizioni di errore durante l'elaborazione"
    return "Script Groovy per elaborazione personalizzata dei dati"

def parse_scripts(files: dict[str, bytes]) -> list[dict]:
    scripts = []
    for name, content in files.items():
        if not name.lower().endswith(".groovy"):
            continue
        source = content.decode("utf-8", errors="replace")
        functions = re.findall(r"\bdef\s+([A-Za-z_][A-Za-z0-9_]*)\s*\(", source)
        strings = re.findall(r'"([^"]{4,})"', source)[:20]
        scripts.append({
            "path": name,
            "name": Path(name).name,
            "sha256": sha256_bytes(content),
            "functions": unique_list(functions),
            "line_count": len(source.splitlines()),
            "source": source,
            "purpose": guess_script_purpose(name, functions, strings),
            "significant_strings": strings,
        })
    return scripts

# ── Mapping parser ───────────────────────────────────────────────────────
def parse_mappings(files: dict[str, bytes]) -> list[dict]:
    mappings = []
    for name, content in files.items():
        if not name.lower().endswith(".mmap"):
            continue
        entry = {
            "path": name,
            "name": Path(name).name,
            "sha256": sha256_bytes(content),
            "entries": [],
            "warnings": [],
        }
        try:
            root = ET.fromstring(content)
        except ET.ParseError:
            entry["warnings"].append("File .mmap non interpretabile come XML")
            mappings.append(entry)
            continue
        destination_bricks = [
            brick
            for brick in root.iter()
            if local_name(brick.tag) == "brick" and brick.attrib.get("type") == "Dst"
        ]
        for destination in destination_bricks:
            target_path = destination.attrib.get("path", "")
            if not target_path:
                continue
            sources = []
            functions = []
            constants = []
            transformation = "Direct"
            details = ""

            def extract_from_node(node):
                nonlocal transformation, details
                for child in node:
                    tag = local_name(child.tag)
                    if tag == "brick":
                        brick_type = child.attrib.get("type", "")
                        brick_path = child.attrib.get("path", "")
                        function_name = child.attrib.get("fname", "")
                        if brick_type == "Src":
                            sources.append(brick_path)
                        elif function_name == "constant":
                            constant_value = brick_path or ""
                            if not constant_value:
                                for grandchild in child:
                                    if local_name(grandchild.tag) == "brick":
                                        constant_value = clean_text(grandchild.attrib.get("path", ""))
                                        break
                            constants.append(constant_value)
                            sources.append(f"[Costante: {constant_value}]")
                        elif function_name == "valuemap":
                            transformation = "Value Mapping"
                            details = "Conversione codici (valuemap)"
                            functions.append("valuemap")
                        elif function_name == "concat":
                            transformation = "UDF: concat"
                            details = "Concatenazione stringhe"
                            functions.append("concat")
                        elif function_name == "iF":
                            transformation = "UDF: iF"
                            details = "Condizione logica"
                            functions.append("iF")
                        elif function_name == "stringEquals":
                            transformation = "UDF: stringEquals"
                            details = "Confronto stringhe"
                            functions.append("stringEquals")
                        elif function_name:
                            transformation = f"UDF: {function_name}"
                            details = f"Funzione custom: {function_name}"
                            functions.append(function_name)
                        if len(list(child)) > 0:
                            extract_from_node(child)
                    elif tag == "arg":
                        extract_from_node(child)

            for element in list(destination):
                if local_name(element.tag) == "arg":
                    extract_from_node(element)
            target_short = target_path.split("/")[-1] if "/" in target_path else target_path
            if constants:
                rule = f"Valore fisso: {', '.join(constants)}"
            elif functions:
                rule = f"Funzione: {', '.join(functions)}"
            else:
                rule = "Copia diretta"
            entry["entries"].append({
                "target": target_path,
                "target_short": target_short,
                "sources": sources,
                "functions": functions,
                "constants": constants,
                "transformation": transformation,
                "regola": rule,
                "details": details,
            })
        if not entry["entries"]:
            entry["warnings"].append("Nessuna associazione estratta dal .mmap")
        else:
            total = len(entry["entries"])
            entry["stats"] = {
                "total": total,
                "direct": sum(1 for item in entry["entries"] if item["transformation"] == "Direct"),
                "constants": sum(1 for item in entry["entries"] if item["constants"]),
                "functions": sum(1 for item in entry["entries"] if item["functions"]),
                "value_mapping": sum(1 for item in entry["entries"] if item["transformation"] == "Value Mapping"),
            }
        mappings.append(entry)
    return mappings

# ── Resource usage detector ──────────────────────────────────────────────
def detect_resource_usage(
    resources: list[dict],
    processes: list[dict],
    message_flows: list[dict],
    participants: list[dict],
) -> list[dict]:
    search_texts: list[tuple[str, str]] = []
    for process in processes:
        search_texts.append((f"Processo: {process.get('name') or process.get('id')}", json.dumps(process, ensure_ascii=False).lower()))
        for step in process.get("steps", []):
            step_name = step.get("name") or step.get("id")
            search_texts.append((f"Step: {step_name}", json.dumps(step, ensure_ascii=False).lower()))
    for message_flow in message_flows:
        search_texts.append((f"Message Flow: {message_flow.get('name') or message_flow.get('id')}", json.dumps(message_flow, ensure_ascii=False).lower()))
    for participant in participants:
        search_texts.append((f"Participant: {participant.get('name') or participant.get('id')}", json.dumps(participant, ensure_ascii=False).lower()))
    for resource in resources:
        resource_name = resource["name"].lower()
        resource_stem = Path(resource["name"]).stem.lower()
        used_by = []
        is_used = False
        for location, text in search_texts:
            if resource_name in text or resource_stem in text:
                is_used = True
                used_by.append(location)
        resource["used"] = is_used
        resource["used_by"] = unique_list(used_by) if used_by else ["-"]
    return resources

# ── Sezione cliente ──────────────────────────────────────────────────────
def describe_adapter(adapter_type: str) -> str:
    adapter_type = adapter_type.lower()
    if "http" in adapter_type:
        return "servizio HTTP/REST"
    if "soap" in adapter_type:
        return "servizio SOAP"
    if "odata" in adapter_type:
        return "servizio OData"
    if "idoc" in adapter_type:
        return "IDoc"
    if "sftp" in adapter_type or "ftp" in adapter_type:
        return "trasferimento file SFTP/FTP"
    if "mail" in adapter_type:
        return "posta elettronica"
    if "jms" in adapter_type:
        return "coda messaggi JMS"
    if "processdirect" in adapter_type:
        return "chiamata diretta a un altro processo CPI"
    if "rfc" in adapter_type:
        return "chiamata RFC SAP"
    return adapter_type or "canale di comunicazione"

def describe_step_type(step: dict) -> str:
    component_type = clean_text(step.get("component_type"))
    labels = {
        "content modifier": "Content Modifier",
        "script groovy": "Script Groovy",
        "message mapping": "Message Mapping",
        "trasformazione XSLT": "Trasformazione XSLT",
        "request reply": "Request Reply",
        "invio messaggio": "Invio messaggio",
        "ricezione messaggio": "Ricezione messaggio",
        "router": "Router / decisione",
        "splitter": "Splitter",
        "gather": "Gather / aggregazione",
        "multicast": "Multicast",
        "content enricher": "Content Enricher",
        "process call": "Chiamata a processo",
        "sottoprocesso": "Sottoprocesso",
        "chiamata a un sottoprocesso": "Chiamata a un sottoprocesso",
        "attività di integrazione CPI": "Attività di integrazione CPI",
        "idoc": "Adapter IDoc",
        "odata": "Adapter OData",
        "soap": "Adapter SOAP",
        "http": "Adapter HTTP/REST",
        "sftp": "Adapter SFTP",
        "jms": "Adapter JMS",
        "mail": "Adapter Mail",
    }
    if component_type in labels:
        return labels[component_type]
    bpmn_type = clean_text(step.get("bpmn_type"))
    bpmn_labels = {
        "startEvent": "Avvio del flusso",
        "endEvent": "Termine del flusso",
        "intermediateCatchEvent": "Evento intermedio",
        "boundaryEvent": "Evento di confine",
        "serviceTask": "Operazione di integrazione",
        "scriptTask": "Esecuzione script",
        "receiveTask": "Ricezione messaggio",
        "sendTask": "Invio messaggio",
        "exclusiveGateway": "Decisione / diramazione condizionale",
        "inclusiveGateway": "Decisione / diramazione condizionale",
        "parallelGateway": "Diramazione / sincronizzazione parallela",
        "callActivity": "Attività di integrazione CPI",
        "subProcess": "Sottoprocesso",
    }
    return bpmn_labels.get(bpmn_type, bpmn_type or "Step BPMN")

def build_client_section(model: dict) -> list[str]:
    paragraphs = []
    name = model["iflow"]["bundle_name"] or model["iflow"]["artifact_id"] or "L'integration flow"
    participants = model.get("participants", [])
    message_flows = model.get("message_flows", [])
    scripts = model.get("scripts", [])
    mappings = model.get("mappings", [])
    xslt_list = model.get("xslt_resources", [])
    routers = model.get("routers", [])
    senders = [p["name"] for p in participants if "sender" in (p.get("type") or "").lower()]
    receivers = [p["name"] for p in participants if "receiv" in (p.get("type") or "").lower()]
    sender_ids = [p.get("id", "") for p in participants if "sender" in (p.get("type") or "").lower()]
    sender_flows = [flow for flow in message_flows if flow.get("source_ref") in sender_ids]
    if senders:
        adapter = describe_adapter(sender_flows[0].get("adapter_type", "")) if sender_flows else "canale di ingresso"
        paragraphs.append(f'**Ricezione**: l\'interfaccia "{name}" riceve i dati da **{", ".join(senders)}** tramite {adapter}.')
    else:
        paragraphs.append(f'**Ricezione**: l\'interfaccia "{name}" acquisisce i dati tramite il canale di ingresso configurato.')
    checks = [s["purpose"] for s in scripts if "check" in s["name"].lower() or "valid" in s["name"].lower() or "Controlla" in s.get("purpose", "")]
    if checks:
        paragraphs.append("**Validazione**: prima di procedere, il flusso esegue i seguenti controlli: " + "; ".join(checks[:5]) + ".")
    else:
        paragraphs.append("**Validazione**: il flusso verifica la presenza e la validità dei dati ricevuti prima di avviare l'elaborazione.")
    transformations = []
    for mapping in mappings:
        count = len(mapping.get("entries", []))
        if count:
            transformations.append(f"il message mapping **{mapping['name']}** ({count} regole di trasformazione)")
    for xslt in xslt_list:
        classifications = xslt.get("classifications", [])
        if "cambio_root_tag" in classifications:
            transformations.append("trasformazioni XSLT che adattano la struttura dei dati per SAP")
        elif "rimozione_struttura" in classifications:
            transformations.append("rimozione della struttura multicanale prima della trasformazione")
    if transformations:
        paragraphs.append("**Trasformazione**: i dati vengono convertiti nel formato SAP tramite: " + "; ".join(transformations[:8]) + ".")
    else:
        paragraphs.append("**Trasformazione**: i dati vengono convertiti nel formato richiesto dal sistema destinatario tramite le regole di mapping configurate.")
    enrichments = [s for s in scripts if any(kw in s["name"].lower() for kw in ("build", "map", "convert"))]
    if enrichments:
        paragraphs.append("**Arricchimento**: durante l'elaborazione, logiche personalizzate completano i dati: " + "; ".join(s["purpose"] for s in enrichments[:4]) + ".")
    receiver_ids = [p.get("id", "") for p in participants if "receiv" in (p.get("type") or "").lower()]
    receiver_flows = [flow for flow in message_flows if flow.get("target_ref") in receiver_ids]
    if receivers:
        adapter = describe_adapter(receiver_flows[0].get("adapter_type", "")) if receiver_flows else "canale di uscita"
        address = receiver_flows[0].get("address", "") if receiver_flows else ""
        extra = f" all'indirizzo `{address}`" if address else ""
        paragraphs.append(f"**Invio**: i dati trasformati vengono inviati a **{', '.join(receivers)}** tramite {adapter}{extra}.")
    if routers:
        for router in routers:
            conditions = [route["condition"] for route in router["routes"] if route["condition"] != "Default"]
            if conditions:
                paragraphs.append(f"**Decisioni**: il flusso instrada i messaggi in base a condizioni configurate (es. `{conditions[0][:100]}`) per differenziare il trattamento dei dati.")
                break
    error_scripts = [s for s in scripts if "error" in s["name"].lower() or "remove" in s["name"].lower()]
    if error_scripts:
        paragraphs.append("**Gestione errori**: in caso di anomalia, il flusso attiva le procedure di gestione delle eccezioni, registrando l'errore e, se configurato, inviando una notifica.")
    return paragraphs

# ── DOCX ─────────────────────────────────────────────────────────────────
def set_cell(cell, text: Any, bold: bool = False, size: int = 9):
    cell.text = ""
    run = cell.paragraphs[0].add_run(clean_text(text) or "-")
    run.bold = bold
    run.font.size = Pt(size)

def set_cell_color(cell, color_hex: str):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shading = tc_pr.makeelement(qn("w:shd"), {qn("w:fill"): color_hex, qn("w:val"): "clear"})
    tc_pr.append(shading)

def add_table(doc, headers: list[str], rows: list[list], color: str = "003366"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell(cell, header, True)
        set_cell_color(cell, color)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell(cells[index], value)
    doc.add_paragraph("")

def add_kv(doc, items):
    add_table(doc, ["Campo", "Valore"], [[key, value] for key, value in items])

def add_field(paragraph, instruction: str, placeholder: str = ""):
    run = paragraph.add_run()
    field_char_begin = OxmlElement("w:fldChar")
    field_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    field_char_separate = OxmlElement("w:fldChar")
    field_char_separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    field_char_end = OxmlElement("w:fldChar")
    field_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(field_char_begin)
    run._r.append(instr_text)
    run._r.append(field_char_separate)
    run._r.append(text)
    run._r.append(field_char_end)

def request_word_field_update(doc):
    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

def configure_heading_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), "99")
    multi_level_type = OxmlElement("w:multiLevelType")
    multi_level_type.set(qn("w:val"), "multilevel")
    abstract_num.append(multi_level_type)
    levels = [("0", "%1.", "1"), ("1", "%1.%2.", "1"), ("2", "%1.%2.%3.", "1")]
    for level_number, level_text, left_indent in levels:
        level = OxmlElement("w:lvl")
        level.set(qn("w:ilvl"), level_number)
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        level.append(start)
        num_format = OxmlElement("w:numFmt")
        num_format.set(qn("w:val"), "decimal")
        level.append(num_format)
        level_text_element = OxmlElement("w:lvlText")
        level_text_element.set(qn("w:val"), level_text)
        level.append(level_text_element)
        level_jc = OxmlElement("w:lvlJc")
        level_jc.set(qn("w:val"), "left")
        level.append(level_jc)
        paragraph_properties = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab_stop = OxmlElement("w:tab")
        tab_stop.set(qn("w:val"), "num")
        tab_stop.set(qn("w:pos"), str((int(level_number) + 1) * 720))
        tabs.append(tab_stop)
        paragraph_properties.append(tabs)
        indentation = OxmlElement("w:ind")
        indentation.set(qn("w:left"), str((int(level_number) + 1) * 720))
        indentation.set(qn("w:hanging"), "360")
        paragraph_properties.append(indentation)
        level.append(paragraph_properties)
        abstract_num.append(level)
    numbering.append(abstract_num)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), "99")
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), "99")
    num.append(abstract_num_id)
    numbering.append(num)
    for level, style_name in enumerate(("Heading 1", "Heading 2", "Heading 3")):
        style = doc.styles[style_name]
        style_ppr = style.element.get_or_add_pPr()
        num_pr = style_ppr.find(qn("w:numPr"))
        if num_pr is None:
            num_pr = OxmlElement("w:numPr")
            style_ppr.append(num_pr)
        ilvl = num_pr.find(qn("w:ilvl"))
        if ilvl is None:
            ilvl = OxmlElement("w:ilvl")
            num_pr.append(ilvl)
        ilvl.set(qn("w:val"), str(level))
        num_id = num_pr.find(qn("w:numId"))
        if num_id is None:
            num_id = OxmlElement("w:numId")
            num_pr.append(num_id)
        num_id.set(qn("w:val"), "99")

def add_cover_page(doc, model: dict):
    interface_name = model["iflow"]["bundle_name"] or model["iflow"]["artifact_id"] or model["source"]["zip_file"]
    version = model["iflow"].get("version") or "-"
    generated_at = model.get("generated_at", "")
    try:
        generated_date = datetime.fromisoformat(generated_at.replace("Z", "+00:00")).strftime("%d/%m/%Y")
    except ValueError:
        generated_date = generated_at or "-"
    document_title = doc.add_paragraph()
    document_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = document_title.add_run("SAP CLOUD INTEGRATION")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    doc.add_paragraph("")
    document_subtitle = doc.add_paragraph()
    document_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = document_subtitle.add_run("DOCUMENTAZIONE INTERFACCIA")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    doc.add_paragraph("")
    doc.add_paragraph("")
    interface_paragraph = doc.add_paragraph()
    interface_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = interface_paragraph.add_run(interface_name)
    run.bold = True
    run.font.size = Pt(20)
    direction_paragraph = doc.add_paragraph()
    direction_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = direction_paragraph.add_run("CENTRIC → SAP")
    run.italic = True
    run.font.size = Pt(14)
    doc.add_paragraph("")
    doc.add_paragraph("─" * 45)
    doc.add_paragraph("")
    for label, value in [("Versione", version), ("Generato il", generated_date), ("Sorgente", model["source"]["zip_file"])]:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label_run = paragraph.add_run(f"{label}: ")
        label_run.bold = True
        paragraph.add_run(clean_text(value) or "-")
    doc.add_paragraph("")
    doc.add_paragraph("")
    footer_text = doc.add_paragraph()
    footer_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_text.add_run("Documentazione tecnica")
    run.italic = True
    run.font.size = Pt(10)

def add_table_of_contents(doc):
    doc.add_page_break()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("INDICE")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    doc.add_paragraph("")
    toc_paragraph = doc.add_paragraph()
    add_field(toc_paragraph, 'TOC \\o "1-3" \\h \\z \\u', "Aggiorna l'indice in Word con Ctrl+A e F9.")
    doc.add_page_break()

def get_participant_role(participant: dict) -> str:
    participant_type = clean_text(participant.get("type")).lower()
    if "sender" in participant_type:
        return "source"
    if "receiv" in participant_type:
        return "target"
    return "other"

def get_process_summary(model: dict) -> dict:
    participants = model.get("participants", [])
    message_flows = model.get("message_flows", [])
    processes = model.get("processes", [])
    mappings = model.get("mappings", [])
    scripts = model.get("scripts", [])
    routers = model.get("routers", [])
    senders = [p.get("name") or p.get("id") for p in participants if get_participant_role(p) == "source"]
    receivers = [p.get("name") or p.get("id") for p in participants if get_participant_role(p) == "target"]
    integration_process = next(
        (p for p in processes if normalize_process_name(p.get("name")) == "integrationprocess"),
        processes[0] if processes else {},
    )
    ordered_processes = order_processes(processes)
    subprocesses = [p for p in ordered_processes if p.get("id") != integration_process.get("id")]
    input_adapter = "-"
    output_adapter = "-"
    sender_ids = {p.get("id") for p in participants if get_participant_role(p) == "source"}
    receiver_ids = {p.get("id") for p in participants if get_participant_role(p) == "target"}
    for flow in message_flows:
        if flow.get("source_ref") in sender_ids and input_adapter == "-":
            input_adapter = flow.get("adapter_type") or "-"
        if flow.get("target_ref") in receiver_ids and output_adapter == "-":
            output_adapter = flow.get("adapter_type") or "-"
    if input_adapter == "-" and message_flows:
        input_adapter = message_flows[0].get("adapter_type") or "-"
    if output_adapter == "-" and message_flows:
        output_adapter = message_flows[-1].get("adapter_type") or "-"
    has_error_handling = any(
        "error" in clean_text(step.get("name")).lower()
        or "exception" in clean_text(step.get("name")).lower()
        or "error" in clean_text(step.get("bpmn_type")).lower()
        for process in processes
        for step in process.get("steps", [])
    )
    if not has_error_handling:
        has_error_handling = any(
            "error" in clean_text(s.get("name")).lower() or "exception" in clean_text(s.get("name")).lower()
            for s in scripts
        )
    elaboration_parts = []
    if mappings:
        elaboration_parts.append("Mapping")
    if scripts:
        elaboration_parts.append("Groovy")
    if model.get("xslt_resources"):
        elaboration_parts.append("XSLT")
    if not elaboration_parts:
        elaboration_parts.append("Trasformazioni configurate")
    return {
        "interface_name": model["iflow"].get("bundle_name") or model["iflow"].get("artifact_id") or "-",
        "sources": ", ".join(senders) or "-",
        "targets": ", ".join(receivers) or "-",
        "input_adapter": describe_adapter(input_adapter),
        "output_adapter": describe_adapter(output_adapter),
        "input_adapter_raw": input_adapter,
        "output_adapter_raw": output_adapter,
        "elaboration": " + ".join(elaboration_parts),
        "has_routing": bool(routers),
        "has_error_handling": has_error_handling,
        "subprocesses": subprocesses,
        "mappings_count": len(mappings),
        "scripts_count": len(scripts),
        "integration_process": integration_process,
    }

def add_executive_summary(doc, model: dict):
    summary = get_process_summary(model)
    doc.add_heading("Sintesi tecnica", level=1)
    summary_rows = [
        ["Interfaccia", summary["interface_name"]],
        ["Sistema sorgente", summary["sources"]],
        ["Sistema destinatario", summary["targets"]],
        ["Protocollo ingresso", summary["input_adapter"]],
        ["Protocollo uscita", summary["output_adapter"]],
        ["Elaborazione", summary["elaboration"]],
        ["Routing", "Sì" if summary["has_routing"] else "No"],
        ["Gestione errori", "Sì" if summary["has_error_handling"] else "Non rilevata"],
        ["Sottoprocessi", str(len(summary["subprocesses"]))],
        ["Message Mapping", str(summary["mappings_count"])],
        ["Script Groovy", str(summary["scripts_count"])],
    ]
    add_table(doc, ["Informazione", "Valore"], summary_rows)

def _add_flow_diagram_legacy(doc, model: dict):
    participants = model.get("participants", [])
    message_flows = model.get("message_flows", [])
    processes = model.get("processes", [])
    routers = model.get("routers", [])

    def participant_label(participant: dict, fallback: str) -> str:
        return clean_text(participant.get("name")) or clean_text(participant.get("id")) or fallback

    def get_role_participants() -> tuple[list[dict], list[dict]]:
        participant_by_id = {p.get("id"): p for p in participants if p.get("id")}
        process_ids = {p.get("id") for p in processes if p.get("id")}
        sender_ids = set()
        receiver_ids = set()
        for p in participants:
            pid = p.get("id")
            pt = clean_text(p.get("type")).lower()
            if "sender" in pt:
                sender_ids.add(pid)
            if "receiver" in pt or "receiv" in pt:
                receiver_ids.add(pid)
        for flow in message_flows:
            src = flow.get("source_ref")
            tgt = flow.get("target_ref")
            if src in participant_by_id and tgt in process_ids:
                sender_ids.add(src)
            if src in process_ids and tgt in participant_by_id:
                receiver_ids.add(tgt)
        for p in participants:
            pid = p.get("id")
            pname = participant_label(p, "").lower()
            if any(w in pname for w in ("centric", "source", "sender", "postgres", "client")):
                sender_ids.add(pid)
            if any(w in pname for w in ("sap", "receiver", "target", "destination", "s4")):
                receiver_ids.add(pid)
        pids = list(participant_by_id.keys())
        if not sender_ids and pids:
            sender_ids.add(pids[0])
        if not receiver_ids and len(pids) >= 2:
            for pid in reversed(pids):
                if pid not in sender_ids:
                    receiver_ids.add(pid)
                    break
        receiver_ids -= sender_ids
        if not receiver_ids and len(pids) >= 2:
            for pid in reversed(pids):
                if pid not in sender_ids:
                    receiver_ids.add(pid)
                    break
        return (
            [participant_by_id[pid] for pid in sender_ids if pid in participant_by_id],
            [participant_by_id[pid] for pid in receiver_ids if pid in participant_by_id],
        )

    def wrap_text(value: str, max_chars: int = 27) -> list[str]:
        words = clean_text(value).split()
        if not words:
            return ["-"]
        lines = []
        current_line = ""
        for word in words:
            candidate = f"{current_line} {word}".strip()
            if len(candidate) <= max_chars:
                current_line = candidate
            else:
                if current_line:
                    lines.append(current_line)
                if len(word) > max_chars:
                    lines.append(word[:max_chars - 3] + "...")
                    current_line = ""
                else:
                    current_line = word
        if current_line:
            lines.append(current_line)
        return lines[:3]

    def get_adapter_labels(sp: list[dict], rp: list[dict]) -> tuple[str, str]:
        sids = {p.get("id") for p in sp if p.get("id")}
        rids = {p.get("id") for p in rp if p.get("id")}
        ia = ""
        oa = ""
        for flow in message_flows:
            adapter = clean_text(flow.get("adapter_type"))
            if not adapter:
                continue
            if flow.get("source_ref") in sids and not ia:
                ia = adapter
            if flow.get("target_ref") in rids and not oa:
                oa = adapter
        return ia, oa

    sender_participants, receiver_participants = get_role_participants()
    source_name = ", ".join(participant_label(p, "Sistema sorgente") for p in sender_participants) or "Sistema sorgente"
    target_name = ", ".join(participant_label(p, "Sistema destinatario") for p in receiver_participants) or "Sistema destinatario"
    input_adapter, output_adapter = get_adapter_labels(sender_participants, receiver_participants)

    integration_process = next(
        (p for p in processes if normalize_process_name(p.get("name")) == "integrationprocess"),
        processes[0] if processes else {},
    )
    process_name = clean_text(integration_process.get("name")) or clean_text(integration_process.get("id")) or "Integration Process"

    important_steps = []
    for step in order_process_steps(integration_process):
        bpmn_type = clean_text(step.get("bpmn_type")).lower()
        component_type = clean_text(step.get("component_type")).lower()
        step_name = clean_text(step.get("name"))
        if bpmn_type in ("startevent", "endevent", "callactivity"):
            continue
        if bpmn_type in ("servicetask", "scripttask", "receivetask", "sendtask", "exclusivegateway", "inclusivegateway", "parallelgateway") or component_type in ("content modifier", "script groovy", "message mapping", "trasformazione xslt", "router", "request reply", "splitter", "gather", "multicast", "content enricher"):
            important_steps.append(step_name or describe_step_type(step))
    important_steps = unique_list(important_steps)[:6]
    if not important_steps:
        important_steps = ["Elaborazione dati"]

    router_names = [clean_text(r.get("name")) or "Router / Decisione" for r in routers if r.get("process_id") == integration_process.get("id")]
    if router_names and not any("router" in s.lower() or "decision" in s.lower() or "choice" in s.lower() for s in important_steps):
        important_steps.append(router_names[0])

    ordered_processes = order_processes(processes)
    subprocesses = [
        clean_text(p.get("name")) or clean_text(p.get("id")) or "Sottoprocesso"
        for p in ordered_processes
        if p.get("id") != integration_process.get("id")
    ][:4]

    # ── Rendering PNG ────────────────────────────────────────────────────
    width = 1650
    background = "#1E1E1E"
    foreground = "#F4F4F4"
    border_color = "#E8E8E8"
    muted_color = "#B8B8B8"
    try:
        font_title = ImageFont.truetype("C:/Windows/Fonts/consolab.ttf", 27)
        font_box = ImageFont.truetype("C:/Windows/Fonts/consola.ttf", 23)
        font_small = ImageFont.truetype("C:/Windows/Fonts/consola.ttf", 19)
    except OSError:
        font_title = ImageFont.load_default()
        font_box = ImageFont.load_default()
        font_small = ImageFont.load_default()

    main_left = 120
    main_right = 650
    main_center = (main_left + main_right) // 2
    source_top = 95
    source_height = 105
    process_top = 350
    process_height = 130 + len(important_steps) * 66
    process_bottom = process_top + process_height
    trunk_start_y = process_bottom + 15
    branch_start_y = process_bottom + 80
    subprocess_left = 825
    subprocess_width = 610
    subprocess_height = 76
    subprocess_gap = 125
    last_branch_y = branch_start_y + (max(len(subprocesses) - 1, 0) * subprocess_gap)
    target_top = max(process_bottom + 350, last_branch_y + 210)
    target_height = 105
    target_bottom = target_top + target_height
    height = target_bottom + 150

    image = Image.new("RGB", (width, height), background)
    draw = ImageDraw.Draw(image)

    def text_centered(lines: list[str], center_x: int, top_y: int, font, fill=foreground, line_spacing: int = 5):
        current_y = top_y
        for line in lines:
            bbox = draw.textbbox((0, 0), line, font=font)
            text_width = bbox[2] - bbox[0]
            text_height = bbox[3] - bbox[1]
            draw.text((center_x - text_width / 2, current_y), line, font=font, fill=fill)
            current_y += text_height + line_spacing

    def draw_dashed_box(left: int, top: int, right: int, bottom: int, dash: int = 12, gap: int = 8):
        for x in range(left, right, dash + gap):
            draw.line((x, top, min(x + dash, right), top), fill=border_color, width=3)
            draw.line((x, bottom, min(x + dash, right), bottom), fill=border_color, width=3)
        for y in range(top, bottom, dash + gap):
            draw.line((left, y, left, min(y + dash, bottom)), fill=border_color, width=3)
            draw.line((right, y, right, min(y + dash, bottom)), fill=border_color, width=3)

    def draw_arrow_vertical(x: int, y_start: int, y_end: int):
        if y_end <= y_start + 20:
            return
        draw.line((x, y_start, x, y_end - 18), fill=foreground, width=3)
        draw.polygon([(x, y_end), (x - 9, y_end - 16), (x + 9, y_end - 16)], fill=foreground)

    def draw_arrow_right(x_start: int, x_end: int, y: int):
        if x_end <= x_start + 20:
            return
        draw.line((x_start, y, x_end - 18, y), fill=foreground, width=3)
        draw.polygon([(x_end, y), (x_end - 16, y - 8), (x_end - 16, y + 8)], fill=foreground)

    text_centered(["Rappresentazione semplificata del flusso"], width // 2, 18, font_title, muted_color)

    source_left = main_center - 150
    source_right = main_center + 150
    source_bottom = source_top + source_height
    draw_dashed_box(source_left, source_top, source_right, source_bottom)
    source_lines = wrap_text(source_name, 22)
    source_text_y = source_top + (source_height - len(source_lines) * 27) // 2
    text_centered(source_lines, main_center, source_text_y, font_box)
    draw_arrow_vertical(main_center, source_bottom + 8, process_top - 24)
    if input_adapter:
        text_centered(wrap_text(describe_adapter(input_adapter), 24), main_center + 115, source_bottom + 38, font_small, muted_color)

    draw_dashed_box(main_left, process_top, main_right, process_bottom)
    text_centered(wrap_text(process_name, 31), main_center, process_top + 28, font_box)
    current_step_y = process_top + 103
    for index, step_name in enumerate(important_steps):
        text_centered(wrap_text(step_name, 34), main_center, current_step_y, font_small)
        if index < len(important_steps) - 1:
            text_centered(["↓"], main_center, current_step_y + 31, font_small)
        current_step_y += 66

    trunk_x = main_center
    if subprocesses:
        trunk_bottom_y = last_branch_y
        draw.line((trunk_x, trunk_start_y, trunk_x, trunk_bottom_y), fill=foreground, width=3)
        for index, subprocess_name in enumerate(subprocesses):
            branch_y = branch_start_y + index * subprocess_gap
            draw_arrow_right(trunk_x, subprocess_left - 24, branch_y)
            box_top = branch_y - subprocess_height // 2
            box_bottom = branch_y + subprocess_height // 2
            box_right = subprocess_left + subprocess_width
            draw_dashed_box(subprocess_left, box_top, box_right, box_bottom)
            subprocess_lines = wrap_text(subprocess_name, 38)
            subprocess_text_y = box_top + (subprocess_height - len(subprocess_lines) * 24) // 2
            text_centered(subprocess_lines, subprocess_left + subprocess_width // 2, subprocess_text_y, font_small)
        draw_arrow_vertical(trunk_x, trunk_bottom_y + 15, target_top - 28)
    else:
        draw_arrow_vertical(trunk_x, trunk_start_y, target_top - 28)

    target_left = main_center - 150
    target_right = main_center + 150
    draw_dashed_box(target_left, target_top, target_right, target_bottom)
    target_lines = wrap_text(target_name, 22)
    target_text_y = target_top + (target_height - len(target_lines) * 27) // 2
    text_centered(target_lines, main_center, target_text_y, font_box)
    if output_adapter:
        text_centered(wrap_text(describe_adapter(output_adapter), 24), main_center + 115, target_top - 58, font_small, muted_color)

    image_stream = BytesIO()
    image.save(image_stream, format="PNG")
    image_stream.seek(0)
    doc.add_heading("Diagramma del flusso", level=1)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(image_stream, width=Pt(500))
    doc.add_paragraph("Rappresentazione semplificata dei sistemi, del processo principale e dei sottoprocessi richiamati.")

def add_flow_diagram(doc, model: dict):
    """Disegna una vista funzionale moderna e leggibile dell'iFlow."""
    participants = model.get("participants", [])
    message_flows = model.get("message_flows", [])
    processes = model.get("processes", [])
    routers = model.get("routers", [])

    def participant_label(participant: dict, fallback: str) -> str:
        return clean_text(participant.get("name")) or clean_text(participant.get("id")) or fallback

    def wrap_text(value: str, max_chars: int = 32, max_lines: int = 3) -> list[str]:
        words = clean_text(value).split()
        if not words:
            return ["-"]
        lines = []
        current = ""
        for word in words:
            candidate = f"{current} {word}".strip()
            if len(candidate) <= max_chars:
                current = candidate
                continue
            if current:
                lines.append(current)
            if len(word) > max_chars:
                lines.append(word[:max_chars - 3] + "...")
                current = ""
            else:
                current = word
        if current:
            lines.append(current)
        if len(lines) <= max_lines:
            return lines
        lines = lines[:max_lines]
        if not lines[-1].endswith("..."):
            lines[-1] = lines[-1][:max_chars - 3].rstrip() + "..."
        return lines

    def get_role_participants() -> tuple[list[dict], list[dict]]:
        participant_by_id = {p.get("id"): p for p in participants if p.get("id")}
        process_ids = {p.get("id") for p in processes if p.get("id")}
        sender_ids = set()
        receiver_ids = set()
        for participant in participants:
            participant_id = participant.get("id")
            participant_type = clean_text(participant.get("type")).lower()
            if "sender" in participant_type:
                sender_ids.add(participant_id)
            if "receiver" in participant_type or "receiv" in participant_type:
                receiver_ids.add(participant_id)
        for flow in message_flows:
            source_ref = flow.get("source_ref")
            target_ref = flow.get("target_ref")
            if source_ref in participant_by_id and target_ref in process_ids:
                sender_ids.add(source_ref)
            if source_ref in process_ids and target_ref in participant_by_id:
                receiver_ids.add(target_ref)
        for participant in participants:
            participant_id = participant.get("id")
            name = participant_label(participant, "").lower()
            if any(token in name for token in ("centric", "source", "sender", "postgres", "client")):
                sender_ids.add(participant_id)
            if any(token in name for token in ("sap", "receiver", "target", "destination", "s4")):
                receiver_ids.add(participant_id)
        participant_ids = list(participant_by_id)
        if not sender_ids and participant_ids:
            sender_ids.add(participant_ids[0])
        receiver_ids -= sender_ids
        if not receiver_ids and len(participant_ids) > 1:
            fallback_receiver = next(
                (pid for pid in reversed(participant_ids) if pid not in sender_ids),
                None,
            )
            if fallback_receiver:
                receiver_ids.add(fallback_receiver)
        return (
            [participant_by_id[pid] for pid in participant_ids if pid in sender_ids],
            [participant_by_id[pid] for pid in participant_ids if pid in receiver_ids],
        )

    def get_adapter_labels(sender_participants: list[dict], receiver_participants: list[dict]) -> tuple[str, str]:
        sender_ids = {p.get("id") for p in sender_participants if p.get("id")}
        receiver_ids = {p.get("id") for p in receiver_participants if p.get("id")}
        input_adapter = ""
        output_adapter = ""
        for flow in message_flows:
            adapter = clean_text(flow.get("adapter_type"))
            if not adapter:
                continue
            if flow.get("source_ref") in sender_ids and not input_adapter:
                input_adapter = adapter
            if flow.get("target_ref") in receiver_ids and not output_adapter:
                output_adapter = adapter
        if not input_adapter and message_flows:
            input_adapter = clean_text(message_flows[0].get("adapter_type"))
        if not output_adapter and message_flows:
            output_adapter = clean_text(message_flows[-1].get("adapter_type"))
        return input_adapter, output_adapter

    sender_participants, receiver_participants = get_role_participants()
    source_name = ", ".join(participant_label(p, "Sistema sorgente") for p in sender_participants) or "Sistema sorgente"
    target_name = ", ".join(participant_label(p, "Sistema destinatario") for p in receiver_participants) or "Sistema destinatario"
    input_adapter, output_adapter = get_adapter_labels(sender_participants, receiver_participants)

    integration_process = next(
        (p for p in processes if normalize_process_name(p.get("name")) == "integrationprocess"),
        processes[0] if processes else {},
    )
    process_name = clean_text(integration_process.get("name")) or clean_text(integration_process.get("id")) or "Integration Process"
    important_steps = []
    for step in order_process_steps(integration_process):
        bpmn_type = clean_text(step.get("bpmn_type")).lower()
        component_type = clean_text(step.get("component_type")).lower()
        if bpmn_type in ("startevent", "endevent", "callactivity"):
            continue
        relevant_bpmn_types = (
            "servicetask", "scripttask", "receivetask", "sendtask",
            "exclusivegateway", "inclusivegateway", "parallelgateway",
        )
        relevant_component_types = (
            "content modifier", "script groovy", "message mapping", "trasformazione xslt",
            "router", "request reply", "splitter", "gather", "multicast", "content enricher",
        )
        if bpmn_type in relevant_bpmn_types or component_type in relevant_component_types:
            important_steps.append({
                "name": clean_text(step.get("name")) or describe_step_type(step),
                "type": describe_step_type(step),
                "bpmn_type": bpmn_type,
                "component_type": component_type,
            })
    unique_steps = []
    seen_step_names = set()
    for step in important_steps:
        key = step["name"].lower()
        if key not in seen_step_names:
            seen_step_names.add(key)
            unique_steps.append(step)
    important_steps = unique_steps[:7]
    if not important_steps:
        important_steps = [{"name": "Elaborazione dati", "type": "Elaborazione", "bpmn_type": "", "component_type": ""}]

    router_names = [
        clean_text(router.get("name")) or "Router / Decisione"
        for router in routers
        if router.get("process_id") == integration_process.get("id")
    ]
    if router_names and not any(
        any(token in step["name"].lower() for token in ("router", "decision", "choice"))
        for step in important_steps
    ):
        important_steps.append({
            "name": router_names[0],
            "type": "Router / decisione",
            "bpmn_type": "exclusivegateway",
            "component_type": "router",
        })

    subprocesses = [
        clean_text(process.get("name")) or clean_text(process.get("id")) or "Sottoprocesso"
        for process in order_processes(processes)
        if process.get("id") != integration_process.get("id")
    ][:4]

    background = "#F7F9FC"
    text_dark = "#182433"
    text_mid = "#607085"
    text_light = "#8995A5"
    line_color = "#AAB8C8"
    border_color = "#D5DEE9"
    source_fill, source_border, source_accent = "#EAF3FF", "#3E8DDD", "#1677D2"
    process_fill, process_border, process_accent = "#FFFFFF", "#2C79C9", "#1268BA"
    subprocess_fill, subprocess_border, subprocess_accent = "#F7F4FF", "#8B72D8", "#7457C6"
    target_fill, target_border, target_accent = "#ECF8F2", "#42A66B", "#218950"

    def load_font(candidates: list[str], size: int):
        for path in candidates:
            try:
                return ImageFont.truetype(path, size)
            except OSError:
                pass
        return ImageFont.load_default()

    regular_fonts = ["C:/Windows/Fonts/segoeui.ttf", "C:/Windows/Fonts/arial.ttf"]
    bold_fonts = ["C:/Windows/Fonts/segoeuib.ttf", "C:/Windows/Fonts/arialbd.ttf"]
    font_title = load_font(bold_fonts, 34)
    font_subtitle = load_font(regular_fonts, 17)
    font_section = load_font(bold_fonts, 19)
    font_name = load_font(bold_fonts, 22)
    font_step = load_font(regular_fonts, 18)
    font_badge = load_font(bold_fonts, 14)
    font_small = load_font(regular_fonts, 14)

    width = 1850
    process_x, process_y, main_width = 90, 410, 950
    process_header, step_height = 105, 78
    process_height = process_header + len(important_steps) * step_height + 30
    process_right, process_bottom = process_x + main_width, process_y + process_height
    source_w, source_h = 500, 155
    source_x, source_y = process_x + (main_width - source_w) // 2, 135
    target_w, target_h = 500, 155
    target_x = source_x
    subprocess_x, subprocess_w, subprocess_h, subprocess_gap = 1190, 540, 110, 28
    subprocess_start_y = process_y + 100
    subprocess_total_h = len(subprocesses) * subprocess_h + max(0, len(subprocesses) - 1) * subprocess_gap
    target_y = max(process_bottom + 150, subprocess_start_y + subprocess_total_h + 100)
    target_bottom = target_y + target_h
    height = max(1050, target_bottom + 110)
    image = Image.new("RGB", (width, height), background)
    draw = ImageDraw.Draw(image)

    def rounded_box(left, top, right, bottom, fill, outline, radius=20, shadow=True, outline_width=2):
        if shadow:
            draw.rounded_rectangle((left + 7, top + 7, right + 7, bottom + 7), radius=radius, fill="#E1E7EF")
        draw.rounded_rectangle((left, top, right, bottom), radius=radius, fill=fill, outline=outline, width=outline_width)

    def draw_text_center(text, x, y, font, fill):
        bbox = draw.textbbox((0, 0), text, font=font)
        text_width = bbox[2] - bbox[0]
        text_height = bbox[3] - bbox[1]
        draw.text((x - text_width / 2, y - text_height / 2 - bbox[1]), text, font=font, fill=fill)

    def draw_lines_center(lines, x, y, font, fill, gap=4):
        heights = [draw.textbbox((0, 0), line, font=font)[3] - draw.textbbox((0, 0), line, font=font)[1] for line in lines]
        current_y = y - (sum(heights) + (len(lines) - 1) * gap) / 2
        for line, line_height in zip(lines, heights):
            bbox = draw.textbbox((0, 0), line, font=font)
            draw.text((x - (bbox[2] - bbox[0]) / 2, current_y), line, font=font, fill=fill)
            current_y += line_height + gap

    def draw_arrow(x1, y1, x2, y2, color=line_color, width_px=4):
        import math
        draw.line((x1, y1, x2, y2), fill=color, width=width_px)
        angle = math.atan2(y2 - y1, x2 - x1)
        length, half_width = 17, 8
        draw.polygon([
            (x2, y2),
            (x2 - length * math.cos(angle) + half_width * math.sin(angle), y2 - length * math.sin(angle) - half_width * math.cos(angle)),
            (x2 - length * math.cos(angle) - half_width * math.sin(angle), y2 - length * math.sin(angle) + half_width * math.cos(angle)),
        ], fill=color)

    def draw_pill(text, x, y, fill, outline, text_fill, font, padding_x=16, padding_y=8):
        bbox = draw.textbbox((0, 0), text, font=font)
        pill_width = bbox[2] - bbox[0] + padding_x * 2
        pill_height = bbox[3] - bbox[1] + padding_y * 2
        draw.rounded_rectangle((x, y, x + pill_width, y + pill_height), radius=pill_height // 2, fill=fill, outline=outline, width=1)
        draw_text_center(text, x + pill_width / 2, y + pill_height / 2, font, text_fill)
        return pill_width, pill_height

    def step_visual(step: dict) -> tuple[str, str, str]:
        component = step["component_type"]
        bpmn_type = step["bpmn_type"]
        visuals = (
            (("router" in component or "gateway" in bpmn_type), "#FFF3D9", "#D68A00", "ROUTER"),
            ("groovy" in component, "#EAF7EF", "#26945D", "GROOVY"),
            ("mapping" in component, "#EEF1FF", "#5366C6", "MAP"),
            ("xslt" in component, "#F2EEFF", "#7959B8", "XSLT"),
            ("content modifier" in component, "#EEF7FF", "#217BC1", "MOD"),
            ("request reply" in component, "#FFF0F0", "#C85A5A", "REQ"),
            ("splitter" in component, "#FFF4E8", "#D17A20", "SPLIT"),
            ("gather" in component, "#F2F5F8", "#60758D", "GATHER"),
            ("multicast" in component, "#F2F5F8", "#60758D", "MULTI"),
        )
        for matches, fill, accent, code in visuals:
            if matches:
                return fill, accent, code
        return "#F0F5FA", "#52708D", "STEP"

    draw_text_center("Integration Flow", width // 2, 42, font_title, text_dark)
    draw_text_center("Rappresentazione funzionale dell'integrazione", width // 2, 78, font_subtitle, text_mid)

    source_right, source_bottom = source_x + source_w, source_y + source_h
    rounded_box(source_x, source_y, source_right, source_bottom, source_fill, source_border, radius=26)
    draw.rounded_rectangle((source_x + 14, source_y + 15, source_x + 22, source_bottom - 15), radius=4, fill=source_accent)
    draw_text_center("SOURCE SYSTEM", source_x + source_w // 2, source_y + 39, font_section, source_accent)
    draw_lines_center(wrap_text(source_name, 36, 2), source_x + source_w // 2, source_y + 92, font_name, text_dark)
    if input_adapter:
        draw_pill(describe_adapter(input_adapter), source_right + 40, source_y + 60, "#FFFFFF", "#CBD5E1", text_mid, font_badge)

    center_x = process_x + main_width // 2
    draw_arrow(center_x, source_bottom + 15, center_x, process_y - 20)
    rounded_box(process_x, process_y, process_right, process_bottom, process_fill, process_border, radius=28, outline_width=3)
    draw.rounded_rectangle((process_x + 16, process_y + 17, process_x + 25, process_bottom - 17), radius=4, fill=process_accent)
    draw_text_center("INTEGRATION PROCESS", center_x, process_y + 34, font_section, process_accent)
    draw_text_center(process_name, center_x, process_y + 72, font_name, text_dark)
    separator_y = process_y + process_header
    draw.line((process_x + 35, separator_y, process_right - 35, separator_y), fill=border_color, width=2)

    timeline_x, step_left, step_right = process_x + 72, process_x + 105, process_right - 35
    step_start_y = separator_y + 20
    for index, step in enumerate(important_steps):
        row_y = step_start_y + index * step_height
        row_center = row_y + step_height // 2
        if index < len(important_steps) - 1:
            draw.line((timeline_x, row_center + 19, timeline_x, row_center + step_height - 3), fill="#C5D3E1", width=3)
        draw.ellipse((timeline_x - 20, row_center - 20, timeline_x + 20, row_center + 20), fill=process_accent)
        draw_text_center(str(index + 1), timeline_x, row_center, font_badge, "#FFFFFF")
        card_top, card_bottom = row_y + 7, row_y + step_height - 7
        rounded_box(step_left, card_top, step_right, card_bottom, "#F8FBFF", "#E0E8F1", radius=16, shadow=False, outline_width=1)
        visual_fill, visual_accent, visual_code = step_visual(step)
        badge_width, _ = draw_pill(visual_code, step_left + 18, card_top + 17, visual_fill, visual_fill, visual_accent, font_badge, 13, 6)
        type_text = clean_text(step["type"])
        type_width = 0
        if type_text:
            type_bbox = draw.textbbox((0, 0), type_text, font=font_small)
            type_width = type_bbox[2] - type_bbox[0] + 30
            draw_text_center(type_text, step_right - type_width / 2, row_center, font_small, text_light)
        name_left = step_left + 18 + badge_width + 24
        name_right = step_right - type_width - 18
        draw_lines_center(wrap_text(step["name"], 46, 2), (name_left + name_right) / 2, row_center, font_step, text_dark, 3)

    if subprocesses:
        draw_text_center("CALLED SUBPROCESSES", subprocess_x + subprocess_w // 2, subprocess_start_y - 26, font_section, subprocess_accent)
        bus_x = process_right + 70
        process_branch_y = process_y + process_height // 2
        last_branch_y = subprocess_start_y + (len(subprocesses) - 1) * (subprocess_h + subprocess_gap) + subprocess_h // 2
        draw.line((process_right, process_branch_y, bus_x, process_branch_y), fill="#A9B8C8", width=3)
        draw.line((bus_x, min(process_branch_y, subprocess_start_y + subprocess_h // 2), bus_x, max(process_branch_y, last_branch_y)), fill="#A9B8C8", width=3)
        for index, subprocess_name in enumerate(subprocesses):
            box_top = subprocess_start_y + index * (subprocess_h + subprocess_gap)
            box_bottom = box_top + subprocess_h
            box_center_y = box_top + subprocess_h // 2
            draw_arrow(bus_x, box_center_y, subprocess_x - 20, box_center_y, "#A9B8C8", 3)
            rounded_box(subprocess_x, box_top, subprocess_x + subprocess_w, box_bottom, subprocess_fill, subprocess_border, radius=20)
            draw.rounded_rectangle((subprocess_x + 14, box_top + 14, subprocess_x + 22, box_bottom - 14), radius=4, fill=subprocess_accent)
            draw_pill("SUBPROCESS", subprocess_x + 34, box_top + 15, "#E8E1FA", "#E8E1FA", subprocess_accent, font_badge, 12, 5)
            draw_lines_center(wrap_text(subprocess_name, 42, 2), subprocess_x + subprocess_w // 2, box_top + 72, font_step, text_dark, 3)

    target_center_x = target_x + target_w // 2
    draw_arrow(center_x, process_bottom + 20, target_center_x, target_y - 20)
    target_right = target_x + target_w
    rounded_box(target_x, target_y, target_right, target_bottom, target_fill, target_border, radius=26)
    draw.rounded_rectangle((target_x + 14, target_y + 15, target_x + 22, target_bottom - 15), radius=4, fill=target_accent)
    draw_text_center("TARGET SYSTEM", target_center_x, target_y + 39, font_section, target_accent)
    draw_lines_center(wrap_text(target_name, 36, 2), target_center_x, target_y + 92, font_name, text_dark)
    if output_adapter:
        draw_pill(describe_adapter(output_adapter), target_right + 40, target_y + 60, "#FFFFFF", "#CBD5E1", text_mid, font_badge)

    legend_items = [
        (source_accent, "Source"),
        (process_accent, "CPI Processing"),
        (subprocess_accent, "Subprocess"),
        (target_accent, "Target"),
    ]
    measurements = []
    for color, label in legend_items:
        bbox = draw.textbbox((0, 0), label, font=font_small)
        measurements.append((color, label, 18 + 10 + bbox[2] - bbox[0] + 35))
    cursor_x = (width - sum(item_width for _, _, item_width in measurements)) // 2
    legend_y = target_bottom + 55
    for color, label, item_width in measurements:
        draw.rounded_rectangle((cursor_x, legend_y - 6, cursor_x + 18, legend_y + 12), radius=8, fill=color)
        draw.text((cursor_x + 28, legend_y - 9), label, font=font_small, fill=text_mid)
        cursor_x += item_width

    image_stream = BytesIO()
    image.save(image_stream, format="PNG", optimize=True)
    image_stream.seek(0)
    doc.add_heading("Diagramma del flusso", level=1)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(image_stream, width=Pt(500))
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption.add_run(
        "Rappresentazione semplificata dei sistemi, del processo principale e dei sottoprocessi richiamati."
    )
    caption_run.font.size = Pt(9)
    caption_run.font.color.rgb = RGBColor(0x89, 0x95, 0xA5)

def add_bold_markdown_paragraph(doc, text: str):
    paragraph = doc.add_paragraph()
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)
    return paragraph

def update_word_fields(docx_path: Path):
    try:
        import win32com.client
    except ImportError:
        print("  ⚠ Indice non aggiornato automaticamente: installa pywin32 con 'pip install pywin32'.")
        return
    word = None
    document = None
    try:
        absolute_path = str(docx_path.resolve())
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        document = word.Documents.Open(absolute_path, ReadOnly=False, AddToRecentFiles=False)
        document.Fields.Update()
        for toc in document.TablesOfContents:
            toc.Update()
        document.Save()
        print("  ✓ Indice Word aggiornato automaticamente.")
    except Exception as exception:
        print(f"  ⚠ Impossibile aggiornare automaticamente l'indice Word: {exception}")
    finally:
        if document is not None:
            document.Close(SaveChanges=True)
        if word is not None:
            word.Quit()
def add_content_modifier_details(
    doc,
    step: dict,
):
    """
    Genera il dettaglio DOCX di un Content Modifier.

    Per i Body:
    - riepilogo del formato;
    - conteggio completo dei riferimenti dinamici;
    - lista completa dei riferimenti, senza taglio;
    - contenuto Body completo fino a 10.000 caratteri.
    """
    content_modifier = step.get("content_modifier")

    if not content_modifier:
        return

    entries = content_modifier.get("entries", [])

    if not entries:
        return

    doc.add_heading(
        f"Content Modifier: {step.get('name') or step.get('id')}",
        level=3,
    )

    content_rows = []

    for entry in entries:
        category = entry.get("category", "-")
        value = entry.get("value", "-")

        expression_label = "No"

        if entry.get("is_expression"):
            expression_label = (
                ", ".join(entry.get("expression_types", []))
                or "Sì"
            )

        # Per Header e Property mostra il valore normalmente.
        if category != "Body":
            content_rows.append(
                [
                    category,
                    entry.get("name", "-"),
                    clean_text(value)[:300],
                    expression_label,
                ]
            )
            continue

        # Per Body mostra solo il riepilogo nella tabella.
        analysis = entry.get("body_analysis", {})

        references = unique_list(
            analysis.get("expression_matches", [])
            or entry.get("expression_matches", [])
        )

        body_summary = [
            f"Formato: {analysis.get('format', '-')}",
            f"Lunghezza: {analysis.get('length', 0)} caratteri",
        ]

        if analysis.get("xml_root"):
            body_summary.append(
                f"Root XML: {analysis['xml_root']} "
                f"({analysis.get('xml_elements', 0)} nodi)"
            )

        if analysis.get("json_keys"):
            body_summary.append(
                "Chiavi JSON: "
                + ", ".join(analysis["json_keys"][:20])
            )

        if references:
            body_summary.append(
                f"Riferimenti rilevati: {len(references)}"
            )
        else:
            body_summary.append("Riferimenti rilevati: 0")

        content_rows.append(
            [
                "Body",
                entry.get("name", "Body"),
                "\n".join(body_summary),
                expression_label,
            ]
        )

    add_table(
        doc,
        [
            "Tipo configurazione",
            "Nome",
            "Valore",
            "Espressione / Mapping",
        ],
        content_rows,
        color="006699",
    )

    # ── Dettaglio esteso dei Body ─────────────────────────────────────────

    for entry in entries:
        if entry.get("category") != "Body":
            continue

        analysis = entry.get("body_analysis", {})
        raw_body = analysis.get("raw_value", "")

        references = unique_list(
            analysis.get("expression_matches", [])
            or entry.get("expression_matches", [])
        )

        if references:
            doc.add_paragraph(
                "Riferimenti dinamici rilevati nel Body:",
                style="Intense Quote",
            )

            # Tutti i riferimenti: nessun [:8] o altro limite.
            reference_rows = [
                [str(index), reference]
                for index, reference in enumerate(references, 1)
            ]

            add_table(
                doc,
                ["#", "Riferimento"],
                reference_rows,
                color="006699",
            )

        if raw_body:
            doc.add_paragraph(
                "Dettaglio Body:",
                style="Intense Quote",
            )

            body_paragraph = doc.add_paragraph()

            # Conserva JSON/XML formattato, senza usare clean_text().
            body_run = body_paragraph.add_run(raw_body[:10000])
            body_run.font.name = "Consolas"
            body_run.font.size = Pt(8)

            if len(raw_body) > 10000:
                doc.add_paragraph(
                    "Il Body è stato troncato nel documento a 10.000 caratteri."
                )

def generate_docx(model: dict, out: Path):
    doc = Document()
    mappings = model.get("mappings", [])
    normal_style = doc.styles["Normal"]
    normal_style.font.size = Pt(10)
    normal_style.font.name = "Calibri"
    for level in ("Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[level]
        style.font.name = "Calibri"
        style.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    doc.styles["Heading 1"].font.size = Pt(14)
    doc.styles["Heading 1"].font.bold = True
    doc.styles["Heading 2"].font.size = Pt(12)
    doc.styles["Heading 2"].font.bold = True
    doc.styles["Heading 3"].font.size = Pt(11)
    doc.styles["Heading 3"].font.bold = True
    configure_heading_numbering(doc)
    request_word_field_update(doc)

    add_cover_page(doc, model)
    add_table_of_contents(doc)
    add_executive_summary(doc, model)
    add_flow_diagram(doc, model)
    doc.add_page_break()

    # ═══ 1. SEZIONE CLIENTE ═══
    doc.add_heading("Cosa fa questa interfaccia", level=1)
    for paragraph in build_client_section(model):
        add_bold_markdown_paragraph(doc, paragraph)
    doc.add_paragraph("")

    # ═══ SISTEMI COINVOLTI ═══
    doc.add_heading("Sistemi coinvolti", level=1)
    participants = model.get("participants", [])
    message_flows = model.get("message_flows", [])
    processes = model.get("processes", [])
    participant_by_id = {p.get("id"): p for p in participants if p.get("id")}
    process_ids = {p.get("id") for p in processes if p.get("id")}
    process_names = {normalize_process_name(p.get("name") or p.get("id")) for p in processes}

    def participant_name(p: dict) -> str:
        return clean_text(p.get("name")) or clean_text(p.get("id")) or "-"

    def is_internal_process_participant(p: dict) -> bool:
        pid = clean_text(p.get("id"))
        pref = clean_text(p.get("process_ref"))
        nname = normalize_process_name(participant_name(p))
        if pref and pref in process_ids:
            return True
        if pid in process_ids:
            return True
        if nname in process_names:
            return True
        return False

    def infer_role(pid: str, flow: dict) -> str:
        p = participant_by_id.get(pid, {})
        pt = clean_text(p.get("type")).lower()
        if "sender" in pt:
            return "Mittente"
        if "receiver" in pt or "receiv" in pt:
            return "Destinatario"
        if flow.get("source_ref") == pid:
            return "Mittente"
        if flow.get("target_ref") == pid:
            return "Destinatario"
        return "Sistema collegato"

    external_systems = []
    seen = set()
    for flow in message_flows:
        adapter_type = clean_text(flow.get("adapter_type"))
        address = clean_text(flow.get("address")) or "-"
        for pid in (flow.get("source_ref"), flow.get("target_ref")):
            p = participant_by_id.get(pid)
            if not p:
                continue
            if is_internal_process_participant(p):
                continue
            system_name = participant_name(p)
            role = infer_role(pid, flow)
            unique_key = (pid, role, adapter_type, address)
            if unique_key in seen:
                continue
            seen.add(unique_key)
            external_systems.append([system_name, role, adapter_type or "Canale configurato", address])
    if external_systems:
        add_table(doc, ["Sistema", "Ruolo", "Adapter", "Indirizzo/Endpoint"], external_systems)
    else:
        doc.add_paragraph("Nessun sistema esterno rilevato dai participant e dai Message Flow.")

    # ═══ ADAPTER ═══
    doc.add_heading("Canali di comunicazione", level=1)
    if message_flows:
        rows = [[flow["name"], flow["adapter_type"], flow["address"] or "-", f"{flow['source_ref']} → {flow['target_ref']}"] for flow in message_flows]
        add_table(doc, ["Nome", "Adapter", "Indirizzo/Endpoint", "Collegamento"], rows)
    else:
        doc.add_paragraph("Nessun message flow rilevato.")

    # ═══ INVENTARIO MAPPING ═══
    if mappings:
        doc.add_heading("Inventario completo dei mapping (formato tecnico)", level=1)
        doc.add_paragraph("Tabella tecnica con tutti i path completi, equivalente al formato CITS.")
        for mapping in mappings:
            if not mapping["entries"]:
                continue
            doc.add_heading(f"Inventario: {mapping['name']}", level=2)
            rows = []
            for index, entry in enumerate(mapping["entries"], 1):
                source = ", ".join(entry.get("sources", ["-"]))
                rows.append([str(index), source, entry.get("transformation", "Direct"), entry.get("target", "-")])
            add_table(doc, ["#", "Source Fields", "Transformation", "Target Field (path completo)"], rows)

    # ═══ XSLT ═══
    xslt_resources = model.get("xslt_resources", [])
    if xslt_resources:
        doc.add_heading("Trasformazioni XSLT", level=1)
        rows = [[xslt["name"], ", ".join(xslt.get("classifications", []))] for xslt in xslt_resources]
        add_table(doc, ["File", "Classificazione"], rows)

    # ═══ GROOVY ═══
    doc.add_heading("Script Groovy", level=1)
    scripts = model["scripts"]
    if scripts:
        for script in scripts:
            doc.add_heading(script["name"], level=2)
            doc.add_paragraph(f"Scopo: {script.get('purpose', '-')}")
            add_kv(doc, [("File", script["path"]), ("Righe", script["line_count"]), ("Funzioni", ", ".join(script.get("functions", [])) or "-")])
    else:
        doc.add_paragraph("Nessuno script Groovy.")

    # ═══ PROCESSI ═══
    doc.add_heading("Processi di integrazione", level=1)
    for process in order_processes(model["processes"]):
        doc.add_heading(process["name"] or process["id"], level=2)
        doc.add_paragraph(f"Step: {len(process['steps'])}  |  Connessioni: {len(process['connections'])}")
        ordered_steps = order_process_steps(process)
        if ordered_steps:
            rows = []
            for index, step in enumerate(ordered_steps, 1):
                rows.append([index, step.get("name") or "-", describe_step_type(step), step["id"]])
            add_table(doc, ["#", "Nome", "Tipo", "ID"], rows)
            for step in ordered_steps:
                add_content_modifier_details(doc, step)
        else:
            doc.add_paragraph("Nessuno step rilevato.")

    # ═══ ROUTER ═══
    routers = model["routers"]
    if routers:
        doc.add_heading("Decisioni e diramazioni", level=1)
        for router in routers:
            doc.add_heading(router["name"] or router["id"], level=2)
            rows = [[route.get("name", "-"), route.get("condition", "Default"), route.get("target_ref", "-")] for route in router["routes"]]
            add_table(doc, ["Ramo", "Condizione", "Destinazione"], rows)

    # ═══ PARAMETRI ═══
    doc.add_heading("Parametri di configurazione", level=1)
    parameters = model["externalized_parameters"]
    if parameters:
        rows = []
        for parameter in parameters:
            used_by = "; ".join(usage["object_name"] for usage in parameter.get("used_by", [])) or "-"
            rows.append([parameter["name"], parameter["type"], parameter.get("value") or parameter.get("default_value", "-"), used_by])
        add_table(doc, ["Parametro", "Tipo", "Valore", "Utilizzato da"], rows)
    else:
        doc.add_paragraph("Nessun parametro.")

    # ═══ RISORSE ═══
    doc.add_heading("Inventario risorse", level=1)
    resources = model["resources"]
    if resources:
        rows = []
        for resource in resources:
            status = "✓ Usata" if resource.get("used") else "✗ Non referenziata"
            used_by = "; ".join(resource.get("used_by", ["-"]))[:120]
            rows.append([resource["name"], resource["type"], status, used_by, resource["path"]])
        add_table(doc, ["Nome", "Tipo", "Stato", "Usato da (Processo/Step)", "Path nel bundle"], rows)
    else:
        doc.add_paragraph("Nessuna risorsa tecnica rilevata.")

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    update_word_fields(out)

# ── Modello principale ───────────────────────────────────────────────────
def parse_iflow_zip(zp: Path) -> dict:
    warnings = []
    errors = []
    with zipfile.ZipFile(zp, "r") as zip_file:
        names = zip_file.namelist()
        files = {name: zip_file.read(name) for name in names if not name.endswith("/")}
    manifest_text = files.get("META-INF/MANIFEST.MF", b"").decode("utf-8", errors="replace")
    manifest = parse_manifest(manifest_text)
    iflow_file = find_iflow_file(names)
    if not iflow_file:
        raise RuntimeError("Nessun file .iflw nel bundle")
    root = ET.fromstring(files[iflow_file])
    participants = parse_participants(root)
    message_flows = parse_message_flows(root)
    processes = parse_processes(root)
    routers = find_router_details(processes)
    parameters = parse_externalized_parameters(files)
    all_objects = []
    for mf in message_flows:
        all_objects.append(("Message Flow", mf.get("name") or mf.get("id"), mf.get("properties", {})))
    for process in processes:
        for step in process["steps"]:
            all_objects.append(("Step", step.get("name") or step.get("id"), step.get("properties", {})))
    for parameter in parameters:
        parameter["used_by"] = []
        for object_type, object_name, object_properties in all_objects:
            for property_key, property_value in object_properties.items():
                if parameter["name"].lower() in clean_text(property_value).lower():
                    parameter["used_by"].append({"object_type": object_type, "object_name": object_name, "property": property_key})
    resources = parse_resources(files)
    scripts = parse_scripts(files)
    mappings = parse_mappings(files)
    xslt_resources = parse_xslt(files)
    resources = detect_resource_usage(resources, processes, message_flows, participants)
    zip_stem = zp.stem
    iflow_stem = Path(iflow_file).stem
    if zip_stem.lower() != iflow_stem.lower():
        warnings.append(f"Nome ZIP ({zip_stem}) ≠ file iFlow ({iflow_stem})")
    if not processes:
        warnings.append("Nessun processo BPMN rilevato")
    return {
        "schema_version": "1.3",
        "generated_at": now_iso(),
        "source": {"zip_file": zp.name, "zip_sha256": sha256_bytes(zp.read_bytes()), "iflow_file": iflow_file, "file_count": len(files)},
        "iflow": {
            "artifact_id": iflow_stem,
            "bundle_name": manifest.get("Bundle-Name", ""),
            "bundle_symbolic_name": manifest.get("Bundle-SymbolicName", ""),
            "version": manifest.get("Bundle-Version", ""),
            "bundle_type": manifest.get("SAP-BundleType", ""),
            "runtime_profile": manifest.get("SAP-RuntimeProfile", ""),
        },
        "manifest": manifest,
        "participants": participants,
        "message_flows": message_flows,
        "processes": processes,
        "routers": routers,
        "externalized_parameters": parameters,
        "resources": resources,
        "scripts": scripts,
        "mappings": mappings,
        "xslt_resources": xslt_resources,
        "quality": {"status": "SUCCESS_WITH_WARNINGS" if warnings else "SUCCESS", "warnings": warnings, "errors": errors},
    }

# ── Batch ────────────────────────────────────────────────────────────────
def process_zip(zp: Path) -> dict:
    print(f"\nElaborazione: {zp.name}")
    try:
        model = parse_iflow_zip(zp)
        artifact_name = safe_filename(model["iflow"]["bundle_name"] or model["iflow"]["artifact_id"] or zp.stem)
        DOCS_DIR.mkdir(parents=True, exist_ok=True)
        JSON_DIR.mkdir(parents=True, exist_ok=True)
        json_output = JSON_DIR / f"{artifact_name}_parsed.json"
        docx_output = DOCS_DIR / f"{artifact_name}.docx"
        json_output.write_text(json.dumps(model, ensure_ascii=False, indent=2), encoding="utf-8")
        generate_docx(model, docx_output)
        print(f"  ✓ JSON: {json_output}")
        print(f"  ✓ DOCX: {docx_output}")
        return {"zip": zp.name, "status": model["quality"]["status"], "json": str(json_output), "docx": str(docx_output), "warnings": model["quality"]["warnings"], "error": ""}
    except Exception as exception:
        print(f"  ✗ ERRORE: {exception}")
        return {"zip": zp.name, "status": "FAILED", "json": "", "docx": "", "warnings": [], "error": str(exception)}

def main():
    INPUT_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    zip_files = sorted(INPUT_DIR.glob("*.zip"))
    if not zip_files:
        print(f"Nessun ZIP in {INPUT_DIR.resolve()}")
        return
    print("=" * 60)
    print("SCPI iFlow Documentation Generator v4")
    print("=" * 60)
    print(f"ZIP: {len(zip_files)}")
    results = [process_zip(zip_file) for zip_file in zip_files]
    report = {
        "generated_at": now_iso(),
        "input": str(INPUT_DIR.resolve()),
        "output": str(OUTPUT_DIR.resolve()),
        "total": len(results),
        "success": sum(1 for r in results if r["status"].startswith("SUCCESS")),
        "failed": sum(1 for r in results if r["status"] == "FAILED"),
        "results": results,
    }
    report_path = OUTPUT_DIR / "batch_report.json"
    report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print("\n" + "=" * 60)
    print(f"COMPLETATO — {report['success']}/{report['total']} OK, {report['failed']} errori")
    print(f"Report: {report_path}")
    print("=" * 60)

if __name__ == "__main__":
    main()